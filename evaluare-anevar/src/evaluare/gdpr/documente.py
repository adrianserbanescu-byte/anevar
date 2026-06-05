"""Generează documentele GDPR model (.docx): politică de prelucrare + acord client.

MODELE/DRAFT — de validat de un jurist și de adaptat la cabinet (vezi
docs/gdpr/*.md). Operatorul de date e evaluatorul, nu aplicația.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument

_DISCLAIMER = (
    "MODEL/DRAFT — a se valida de un jurist (GDPR) și a se adapta la cabinet înainte de folosire. "
    "Operatorul de date cu caracter personal este evaluatorul/cabinetul, nu aplicația software."
)


def _disclaimer(doc: DocxDocument) -> None:
    doc.add_paragraph().add_run(_DISCLAIMER).bold = True


def genereaza_politica_gdpr() -> DocxDocument:
    """Politica de prelucrare a datelor (model)."""
    doc = Document()
    doc.add_heading("POLITICĂ DE PRELUCRARE A DATELOR CU CARACTER PERSONAL", level=0)
    _disclaimer(doc)
    sectiuni = [
        ("1. Operator", "[Nume evaluator / Cabinet], CIF/CNP ____, adresă ____, e-mail ____, telefon ____."),
        ("2. Categorii de date", "Date de identificare ale clientului/proprietarului (nume, adresă, "
            "CNP dacă e necesar, contact); date despre imobil (adresă, cadastral, CF, fotografii); "
            "date KYC/beneficiar real (în cazul AML)."),
        ("3. Scop și temei legal", "Întocmirea raportului de evaluare — executarea contractului "
            "(art. 6(1)(b) GDPR) și interes legitim profesional (art. 6(1)(f)). Conformitate AML — "
            "obligație legală (art. 6(1)(c); Legea 129/2019)."),
        ("4. Sub-procesatori (asistent AI)", "Pentru redactarea textului de analiză se poate folosi "
            "un serviciu AI extern (Anthropic/Perplexity), căruia i se transmite DOAR text anonimizat "
            "(nume/adresă/CF/cadastral înlocuite cu marcaje înainte de trimitere). Aplicația poate "
            "funcționa și complet offline (fără AI, fără transfer extern) — recomandat pentru cazuri "
            "sensibile. De clarificat juridic: acord DPA cu furnizorul AI, localizarea serverelor, "
            "politica de retenție."),
        ("5. Perioada de păstrare", "Conform obligațiilor profesionale ANEVAR și AML (de regulă 5 ani "
            "de la încheierea relației)."),
        ("6. Drepturile persoanei vizate", "Acces, rectificare, ștergere, restricționare, opoziție, "
            "portabilitate; plângere la ANSPDCP."),
        ("7. Securitate", "Date păstrate local (calculatorul evaluatorului), bază SQLite, backup "
            "periodic, acces restricționat."),
    ]
    for titlu, text in sectiuni:
        doc.add_heading(titlu, level=1)
        doc.add_paragraph(text)
    return doc


def genereaza_consimtamant_gdpr(client_nume: str = "", adresa: str = "") -> DocxDocument:
    """Acord de prelucrare semnat de client (model). Pre-completează nume/adresă dacă sunt date."""
    doc = Document()
    doc.add_heading("ACORD PRIVIND PRELUCRAREA DATELOR CU CARACTER PERSONAL", level=0)
    _disclaimer(doc)
    doc.add_paragraph(
        f"Subsemnatul(a) {client_nume or '____________________'}, în calitate de client/proprietar, "
        f"în legătură cu evaluarea imobilului situat la {adresa or '____________________'}, declar că "
        "am fost informat(ă) și sunt de acord ca evaluatorul/cabinetul (operator de date) să prelucreze "
        "datele mele cu caracter personal în scopul întocmirii raportului de evaluare și al conformității "
        "legale (inclusiv AML, Legea 129/2019), conform GDPR (Regulamentul UE 2016/679)."
    )
    doc.add_paragraph("Am fost informat(ă) că:")
    for p in (
        "pentru redactarea analizei se poate folosi un asistent software cu componentă AI, căruia i se "
        "transmite doar text anonimizat (fără nume, adresă exactă, CNP, CF);",
        "pot solicita varianta complet offline (fără niciun transfer extern);",
        "am dreptul de acces, rectificare, ștergere, opoziție și de a depune plângere la ANSPDCP;",
        "datele se păstrează conform obligațiilor profesionale și legale aplicabile.",
    ):
        doc.add_paragraph(p, style="List Bullet")
    doc.add_paragraph("☐ Sunt de acord cu folosirea asistentului AI (pe text anonimizat).")
    doc.add_paragraph("☐ Solicit prelucrarea exclusiv offline (fără AI).")
    doc.add_paragraph()
    doc.add_paragraph("Nume și prenume: ____________________   Data: __________   Semnătura: ____________________")
    return doc


def salveaza(doc: DocxDocument, cale: Path | str) -> Path:
    cale = Path(cale)
    doc.save(str(cale))
    return cale
