"""Generator de raport .docx in structura GBF / SEV (shell complet + 7 capitole).

Shell GBF: copertA, scrisoare de transmitere, declaratie de conformitate, termeni de
referintA, cele 7 capitole de analiza (raportare conform SEV 106), alocarea valorii, riscul GEV 520, anexe, semnatura.
"""
from __future__ import annotations

import base64
import binascii
from decimal import Decimal, InvalidOperation
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor

from evaluare.logging_setup import get_logger
from evaluare.models.report_context import ReportContext

log = get_logger(__name__)

# Note de provenienta (mod demo/review): explica ce e calculat, extras, AI, exemplu sau placeholder.
ADNOTARI = {
    "legenda": ("LEGENDA NOTELOR DEMO (text colorat, NU apare in raportul real): "
                "[CALCULAT]=produs de motorul de calcul; [EXTRAS]=preluat automat din anunt; "
                "[INTRODUS]=completat de evaluator; [AI]=text generat de inteligenta artificiala; "
                "[SABLON]=text fix standard; [EXEMPLU]/[PLACEHOLDER]=valori demonstrative de inlocuit."),
    "coperta": ("[INTRODUS] Identificarea (client, cadastral, CF, evaluator) e introdusa de evaluator — "
                "aici valori generice/placeholder. [CALCULAT] VALOAREA estimata e calculata real de motor."),
    "scrisoare": "[SABLON+CALCULAT] Text-sablon fix, completat automat cu valoarea calculata. Functional.",
    "declaratie": "[SABLON] Declaratii standard ANEVAR — text fix. Functional.",
    "termeni": ("[INTRODUS] Compilati automat din campurile lucrarii (client, beneficiar, scop, moneda, "
                "curs, date). Functional."),
    "cap1": "[CALCULAT] Sinteza generata automat din rezultatele calculate. Functional.",
    "cap2": "[AI] Text narativ generat de inteligenta artificiala; de revizuit de evaluator.",
    "cap3": "[AI] Analiza de piata descriptiva (AI) — NU pe o baza de date de tranzactii reale.",
    "cap4": ("[EXTRAS/INTRODUS] Date fizice extrase din anunt / introduse de evaluator (reale). "
             "Nr. cadastral si CF aici sunt PLACEHOLDER."),
    "cap5": "[AI] Analiza celei mai bune utilizari generata de AI.",
    "cap6": ("[CALCULAT] Grilele (teren + casa) si costul CIN sunt calculate real de motor. "
             "[EXTRAS] Preturile/suprafetele comparabilelor sunt reale (din anunturi). "
             "[EXEMPLU] Ajustarile din grila sunt demonstrative — de stabilit de evaluator la inspectie."),
    "cap7": "[CALCULAT] Selectia metodei si valoarea finala sunt calculate. [AI] Textul explicativ.",
    "alocare": "[CALCULAT] Valoare constructii = valoare proprietate − valoare teren.",
    "gev520": "[AI] Text standard GEV 520 generat de AI; de adaptat la caz.",
    "anexe": ("[EXTRAS] Sursele comparabilelor sunt linkuri reale. [EXEMPLU] Fotografiile sunt "
              "exemplificative. [PLACEHOLDER] Documentele cadastrale — de atasat."),
    "semnatura": "[PLACEHOLDER] Caseta de semnatura — de completat de evaluator.",
}


def _nota(doc: DocxDocument, cheie: str, adnotari: bool) -> None:
    """Adauga o nota de provenienta (colorata, italic) sub un titlu — doar in modul demo."""
    if not adnotari:
        return
    text = ADNOTARI.get(cheie, "")
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run("▮ NOTA DEMO — " + text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0xB0, 0x6A, 0x00)


# Disclaimer al APLICATIEI catre evaluator — apare in fruntea ORICARUI document generat (cerinta Adi):
# documentul e un DRAFT, verificarea + raspunderea revin evaluatorului ANEVAR semnatar, iar aplicatia
# nu raspunde de datele introduse / rezultate. Constanta e refolosibila si de alte generatoare.
DISCLAIMER_APLICATIE = (
    "Acest document este un DRAFT generat automat, ca instrument de asistență. Verificarea integrală "
    "(date, conținut, valori) și răspunderea profesională pentru raportul final revin evaluatorului "
    "autorizat ANEVAR care îl semnează. Aplicația nu poartă nicio răspundere pentru datele introduse "
    "sau pentru rezultatele prezentate."
)


def _disclaimer_aplicatie(doc: DocxDocument) -> None:
    """Disclaimer-ul aplicatiei catre evaluator, in fruntea documentului (mereu, nu doar in demo)."""
    p = doc.add_paragraph()
    t = p.add_run("⚠ NOTĂ A APLICAȚIEI CĂTRE EVALUATOR")
    t.bold = True
    t.font.size = Pt(10)
    t.font.color.rgb = RGBColor(0xB0, 0x6A, 0x00)
    r = doc.add_paragraph().add_run(DISCLAIMER_APLICATIE)
    r.italic = True
    r.font.size = Pt(9)


def _narativ(ctx: ReportContext, capitol: str) -> str | None:
    """Returneaza textul narativ pentru un capitol, daca exista."""
    for sectiune in ctx.narrative:
        if sectiune.capitol == capitol:
            return sectiune.text
    return None


def _fmt(v) -> str:
    """Formateaza un numar cu separator de mii (stil RO) fara zecimale inutile."""
    try:
        d = Decimal(str(v)).quantize(Decimal("1"))
        return f"{int(d):,}".replace(",", ".")
    except (InvalidOperation, ValueError, OverflowError, TypeError) as e:
        log.debug("Valoare neformatabilă numeric (%r), folosesc str(): %s", v, e)
        return str(v)


# Denumirea + SURSA/definitia tipului valorii (cerinta SEV 102 §20.4 + SEV 106 §30.6(i):
# raportul citeaza tipul valorii SI sursa definitiei). Transforma slug-ul intern in text cu referinta.
_TIP_VALOARE_TXT = {
    "piata": "valoare de piață (definită conform SEV 102 — Tipuri ale valorii / IVS 104)",
    "justa": "valoare justă (definită conform SEV 102 / IFRS 13)",
    "lichidare": "valoare de lichidare (definită conform SEV 102, Anexă)",
    "asigurare": "valoare de asigurare — cost de înlocuire brut (de reconstrucție), conform SEV 450",
    "impozabila": "valoare impozabilă (estimată conform GEV 500)",
}


def _tip_valoare_txt(t: str) -> str:
    """Denumirea + sursa tipului valorii. Slug cunoscut -> text cu referinta; frază deja citata
    (conține SEV/IVS/IFRS) -> neschimbată; altfel adaugă referinta SEV 102."""
    raw = (t or "").strip()
    key = raw.lower()
    if key in _TIP_VALOARE_TXT:
        return _TIP_VALOARE_TXT[key]
    if any(s in raw for s in ("SEV", "IVS", "IFRS")):
        return raw
    return f"{raw} (conform SEV 102 — Tipuri ale valorii)"


def _b2(v) -> str:
    """Rotunjeste la 2 zecimale (preturi unitare, EUR/mp)."""
    try:
        return f"{Decimal(str(v)).quantize(Decimal('0.01'))}"
    except (InvalidOperation, ValueError, TypeError) as e:
        log.debug("Valoare neformatabilă numeric (%r), folosesc str(): %s", v, e)
        return str(v)


def _pct(v) -> str:
    """Fractie -> procent cu 2 zecimale (0.1727 -> '17.27%')."""
    try:
        return f"{(Decimal(str(v)) * 100).quantize(Decimal('0.01'))}%"
    except (InvalidOperation, ValueError, TypeError) as e:
        log.debug("Valoare neformatabilă numeric (%r), folosesc str(): %s", v, e)
        return str(v)


def _fara_tva(ctx: ReportContext) -> str:
    return "Valoarea nu contine TVA." if ctx.reconciled.valoare_fara_tva else ""


def _echiv_lei(ctx: ReportContext) -> str:
    """Echivalentul in LEI al valorii finale (daca moneda e EUR si exista curs BNR)."""
    curs = ctx.meta.curs_valutar
    if curs is None or (ctx.meta.moneda or "").upper() == "LEI":
        return ""
    lei = ctx.reconciled.valoare_finala * curs
    return f" (echivalent ~{_fmt(lei)} LEI la cursul BNR {curs})"


def _valoare_teren(ctx: ReportContext):
    """Valoarea terenului folosita in raport (din grila daca exista, altfel din alocare)."""
    if ctx.land_result is not None:
        return ctx.land_result.valoare_teren
    if ctx.alocare_constructii is not None:
        return ctx.reconciled.valoare_finala - ctx.alocare_constructii
    return None


# --------------------------------------------------------------------------- #
# Front matter (shell GBF)
# --------------------------------------------------------------------------- #
def _coperta(doc: DocxDocument, ctx: ReportContext, adnotari: bool = False) -> None:
    meta = ctx.meta
    doc.add_heading("RAPORT DE EVALUARE", level=0)
    _nota(doc, "legenda", adnotari)
    _nota(doc, "coperta", adnotari)
    if meta.nr_lucrare:        # numar de identificare a raportului (Procedura §6: recomandat pe coperta)
        doc.add_paragraph(f"Nr. de identificare raport: {meta.nr_lucrare}")
    p = doc.add_paragraph()
    p.add_run("Proprietate imobiliara: casa de locuit si teren aferent").bold = True
    doc.add_paragraph(f"Adresa: {meta.adresa}")
    doc.add_paragraph(f"Numar cadastral: {meta.numar_cadastral}; {meta.carte_funciara}")
    doc.add_paragraph(f"Client: {meta.client_nume} ({meta.client_tip})")
    if meta.proprietar:
        doc.add_paragraph(f"Proprietar: {meta.proprietar}")
    if meta.beneficiar:
        doc.add_paragraph(f"Beneficiar / Finantator: {meta.beneficiar}")
    doc.add_paragraph(f"Scopul evaluarii: {meta.scop}")
    doc.add_paragraph(
        f"Evaluator: {meta.evaluator_nume}, membru ANEVAR, "
        f"legitimatia {meta.evaluator_legitimatie}"
    )
    doc.add_paragraph(
        f"Data evaluarii: {meta.data_evaluarii}; Data raportului: {meta.data_raportului}"
    )
    vp = doc.add_paragraph()
    vp.add_run(
        f"VALOAREA ESTIMATA: {_fmt(ctx.reconciled.valoare_finala)} {meta.moneda}"
        f"{_echiv_lei(ctx)}. {_fara_tva(ctx)}"
    ).bold = True
    doc.add_page_break()


def _scrisoare_transmitere(doc: DocxDocument, ctx: ReportContext, adnotari: bool = False) -> None:
    meta = ctx.meta
    doc.add_heading("SCRISOARE DE TRANSMITERE", level=1)
    _nota(doc, "scrisoare", adnotari)
    catre = meta.client_nume + (f" / {meta.beneficiar}" if meta.beneficiar else "")
    doc.add_paragraph(f"Catre: {catre}")
    doc.add_paragraph(
        f"Va transmitem raportul de evaluare a proprietatii imobiliare situate in "
        f"{meta.adresa}, intocmit in scopul: {meta.scop}. In urma analizelor efectuate, "
        f"valoarea de piata estimata a proprietatii este de {_fmt(ctx.reconciled.valoare_finala)} "
        f"{meta.moneda}, la data evaluarii ({meta.data_evaluarii}). {_fara_tva(ctx)}"
    )
    conformitate = (
        "Raportul a fost elaborat in conformitate cu Standardele de evaluare a bunurilor "
        "ANEVAR, editia 2025 (in vigoare de la 1 iulie 2025, aprobate prin HCN nr. 2/2025)"
    )
    clauza = _ghid_clauza(ctx)
    if clauza:
        conformitate += " si cu Ghidul " + clauza.removeprefix("incluzand Ghidul ")
    conformitate += " Poate fi utilizat exclusiv in scopul mentionat, de catre utilizatorul desemnat."
    doc.add_paragraph(conformitate)
    doc.add_paragraph(
        f"Cu deosebita consideratie, {meta.evaluator_nume}, "
        f"membru ANEVAR, legitimatia {meta.evaluator_legitimatie}."
    )


_GHID_CLAUZA = {
    "GEV_520": "incluzand Ghidul GEV 520 — Evaluarea pentru garantarea imprumutului.",
    "GEV_630": "incluzand Ghidul GEV 630 — Evaluarea bunurilor imobile.",
    "SEV_450": "incluzand SEV 450 — Evaluarea costurilor in scop de asigurare.",
    "GEV_500": "incluzand Ghidul GEV 500 — Estimarea valorii impozabile a cladirilor.",
    "none": "",
}


def _ghid_clauza(ctx: ReportContext) -> str:
    """Clauza de ghid aplicabil (GEV) din profilul evaluarii, pentru declaratia de conformitate."""
    return _GHID_CLAUZA.get(ctx.profil.ghid, "")


def _declaratie_conformitate(doc: DocxDocument, ctx: ReportContext, adnotari: bool = False) -> None:
    doc.add_heading("DECLARATIE DE CONFORMITATE SI CERTIFICARE", level=1)
    _nota(doc, "declaratie", adnotari)
    afirmatii = [
        "Prezentul raport a fost elaborat in conformitate cu Standardele de evaluare a bunurilor "
        "ANEVAR (SEV), editia 2025, in vigoare de la 1 iulie 2025 (HCN nr. 2/2025), "
        + _ghid_clauza(ctx),
        "Faptele prezentate in raport sunt reale si corecte, dupa cunostinta evaluatorului.",
        "Analizele, opiniile si concluziile sunt limitate numai de ipotezele si conditiile "
        "limitative prezentate.",
        "Evaluatorul nu are niciun interes prezent sau de perspectiva in proprietatea evaluata "
        "si niciun interes personal cu privire la partile implicate.",
        "Onorariul evaluatorului nu este conditionat de exprimarea unei valori predeterminate "
        "sau de marimea valorii estimate.",
        "Evaluatorul este membru ANEVAR, detine competenta necesara si asigurarea de raspundere "
        "civila profesionala in vigoare.",
        # SEV 100 (2025) — declarate explicit in raport (le aplicam de-facto prin validari + audit):
        "Evaluatorul a aplicat scepticism profesional pe parcursul evaluarii, evaluand critic datele "
        "si informatiile utilizate si neacceptandu-le necritic (SEV 100, par. 10.4).",
        "Raportul a fost supus unei proceduri de verificare a calitatii (controale de coerenta a "
        "datelor de intrare si a calculelor) inainte de finalizare, conform SEV 100, par. 20.",
    ]
    for a in afirmatii:
        doc.add_paragraph(a, style="List Bullet")


def _termeni_referinta(doc: DocxDocument, ctx: ReportContext, adnotari: bool = False) -> None:
    meta = ctx.meta
    doc.add_heading("TERMENI DE REFERINTA AI EVALUARII", level=1)
    _nota(doc, "termeni", adnotari)
    doc.add_paragraph(f"Clientul evaluarii: {meta.client_nume} ({meta.client_tip}).")
    if meta.beneficiar:
        doc.add_paragraph(f"Beneficiarul / utilizatorul desemnat: {meta.beneficiar}.")
    doc.add_paragraph(f"Scopul evaluarii: {meta.scop}.")
    doc.add_paragraph(f"Tipul valorii estimate: {_tip_valoare_txt(meta.tip_valoare)}.")
    moneda_txt = f"Moneda raportarii: {meta.moneda}."
    if meta.curs_valutar is not None:
        moneda_txt += f" Curs de schimb EUR/LEI: {meta.curs_valutar}."
    doc.add_paragraph(moneda_txt)
    date_txt = f"Data evaluarii: {meta.data_evaluarii}. Data raportului: {meta.data_raportului}."
    if meta.data_inspectiei:
        date_txt += f" Data inspectiei: {meta.data_inspectiei}."
    if meta.valabilitate:
        date_txt += f" Valabilitate: {meta.valabilitate}."
    doc.add_paragraph(date_txt)
    doc.add_paragraph(
        f"Identificarea proprietatii: {meta.adresa}, nr. cadastral {meta.numar_cadastral}, "
        f"{meta.carte_funciara}."
    )
    doc.add_paragraph(f"Dreptul evaluat: {meta.tip_drept} (SEV 230).")
    if meta.act_proprietate:
        doc.add_paragraph(f"Act de proprietate: {meta.act_proprietate}.")
    doc.add_paragraph(
        f"Sarcini / grevari (extras CF): {meta.sarcini}." if meta.sarcini
        else "Sarcini / grevari (extras CF): nu au fost declarate; de verificat în extrasul de carte funciară "
             "actualizat (relevant pentru garantare)."
    )
    premisa = (
        "Premise: proprietatea este evaluata in ipoteza utilizarii continue, libera de sarcini, "
        "cu exceptia celor mentionate explicit in raport."
    )
    if ctx.profil.ghid == "GEV_520":
        premisa += (
            " Daca este ocupata de proprietar, se evalueaza in ipoteza transferului ca bun "
            "liber/disponibil (GEV 520, A8)."
        )
    doc.add_paragraph(premisa)
    if ctx.profil.ghid == "GEV_520":
        # GEV 520 A3 (SEV 101): independenta / implicare materiala
        doc.add_paragraph(
            "Independenta: evaluatorul declara ca nu are nicio implicare materiala, prezenta sau "
            "anterioara, cu bunul evaluat, cu debitorul sau cu un debitor potential, care sa afecteze "
            "obiectivitatea (GEV 520, A3; SEV 101). Utilizatorul desemnat al raportului este creditorul "
            "nominalizat; orice alta utilizare necesita personalizare de catre evaluator."
        )
        # GEV 520 A4: ipoteze speciale (ex. vanzare fortata / perioada de marketing limitata)
        doc.add_paragraph(
            "Ipoteze speciale: daca se solicita o valoare in premisa unei vanzari fortate sau cu "
            "perioada de marketing limitata, aceasta se precizeaza distinct; valoarea pe ipoteza "
            "speciala este valabila numai la data evaluarii si poate sa nu fie realizabila la o data "
            "viitoare (GEV 520, A4-A5)."
        )
    # SEV 101, 20.1 — elemente suplimentare ale termenilor de referinta
    doc.add_paragraph(
        f"Evaluatorul: {meta.evaluator_nume}, membru ANEVAR, legitimatia "
        f"{meta.evaluator_legitimatie} (SEV 101, 20.1 e)."
    )
    doc.add_paragraph(
        "Natura si amploarea activitatilor: inspectia si analiza in limitele mandatului de evaluator "
        "autorizat, fara investigatii distructive, geotehnice sau juridice de specialitate (SEV 101, "
        "20.1 i)."
    )
    doc.add_paragraph(
        "Sursa informatiilor: documentele proprietatii puse la dispozitie, date de piata din oferte "
        "publice (cu caracter indicativ, verificate de evaluator) si cursul de schimb BNR (SEV 101, "
        "20.1 j)."
    )
    doc.add_paragraph(
        "Factori de mediu, sociali si de guvernanta (ESG): se au in vedere factorii cunoscuti, "
        "relevanti pentru valoare (ex. performanta energetica conform certificatului energetic, daca "
        "exista); nu s-au efectuat investigatii de mediu specializate (SEV 101/106, 20.1/30.6 m)."
    )
    doc.add_paragraph(
        "Specialist: nu a fost cazul utilizarii unui specialist sau furnizor extern de servicii "
        "(SEV 106, 30.6 o)."
    )
    doc.add_paragraph(
        "Tipul raportului si restrictii: raport scris, narativ. Raportul poate fi utilizat exclusiv "
        "in scopul declarat, de catre utilizatorul desemnat; difuzarea sau publicarea, integrala ori "
        "partiala, se face numai cu acordul scris al evaluatorului (SEV 101, 20.1 n, o)."
    )
    # Transparenta privind instrumentul software + AI (mitigare a perceptiei la verificarea bancara):
    # numerele sunt DETERMINISTE; AI scrie doar proza; evaluatorul verifica si isi asuma.
    p_ai = doc.add_paragraph()
    p_ai.add_run("Asistare software si componenta AI (transparenta): ").bold = True
    p_ai.add_run(
        "raportul a fost intocmit cu ajutorul unui instrument software de asistenta a evaluarii. "
        "TOATE valorile numerice din raport (costuri, grilele de comparatie si ajustari, "
        "capitalizare / DCF, reconcilierea, alocarea valorii si valoarea finala) sunt DETERMINISTE — "
        "sunt calculate de motorul de evaluare si reproductibile din datele de intrare; ele NU sunt "
        "generate de inteligenta artificiala. Textul narativ explicativ a fost generat cu asistenta "
        "instrumentului software (optional cu asistenta AI, pe date anonimizate, fara transmiterea "
        "datelor cu caracter personal) si a fost verificat, corectat si asumat integral de evaluatorul "
        "autorizat ANEVAR, caruia ii apartine valoarea si responsabilitatea profesionala. Instrumentul "
        "asista evaluatorul; nu il inlocuieste si nu decide valoarea."
    )


# --------------------------------------------------------------------------- #
# Tabele de calcul
# --------------------------------------------------------------------------- #
def _adauga_grila_comparatie(doc: DocxDocument, ctx: ReportContext) -> None:
    if ctx.market_result is None or not ctx.comparables:
        doc.add_paragraph("Abordarea prin piata nu a fost aplicata (a se vedea reconcilierea).")
        return
    doc.add_paragraph("Grila de comparatie directa pe pret total (proprietate intreaga):")
    m = ctx.market_result
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Comparabil"
    hdr[1].text = "Pret total corectat"
    hdr[2].text = "Ajustare bruta"
    hdr[3].text = "În medie"
    for i, _comp in enumerate(ctx.comparables):
        pret = m.preturi_unitare_corectate[i] if i < len(m.preturi_unitare_corectate) else ""
        bruta = m.ajustari_brute[i] if i < len(m.ajustari_brute) else ""
        row = table.add_row().cells
        row[0].text = f"{i + 1}"
        row[1].text = _b2(pret)
        row[2].text = _pct(bruta)
        # #6 (verificare adversariala B): valoarea = MEDIA celor N comparabile -> marcam TOATE cele N
        # randuri din indici_mediati cu „DA" (nu doar unul), ca headline-ul sa se reconcilieze cu grila.
        marcaj = "DA" if i in m.indici_mediati else ""
        if i == m.index_selectat:
            marcaj = (marcaj + " *").strip()   # * = cel mai similar (ajustare bruta minima), pastrat ca referinta
        row[3].text = marcaj
    doc.add_paragraph(
        f"Valoarea prin comparatie de piata: {_fmt(m.valoare_piata)} {ctx.meta.moneda} = media celor mai "
        f"similare {len(m.indici_mediati)} comparabile (cele marcate cu DA; * = cel mai similar, ajustare minima)."
    )
    doc.add_paragraph(
        "Sursele comparabilelor au caracter indicativ (oferte de pe portaluri publice); "
        "preturile si atributele au fost verificate si retinute de evaluator pe baza "
        "rationamentului profesional."
    )


def _adauga_grila_teren(doc: DocxDocument, ctx: ReportContext) -> None:
    if ctx.land_result is None or not ctx.land_comparables:
        return
    doc.add_paragraph("Grila de comparatie pentru teren (EUR/mp):")
    lr = ctx.land_result
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Comparabil teren"
    hdr[1].text = "Pret/mp corectat"
    hdr[2].text = "Ajustare bruta"
    hdr[3].text = "În medie"
    for i, c in enumerate(ctx.land_comparables):
        pret = lr.preturi_mp_corectate[i] if i < len(lr.preturi_mp_corectate) else ""
        bruta = lr.ajustari_brute[i] if i < len(lr.ajustari_brute) else ""
        row = table.add_row().cells
        eticheta = f"{i + 1}" + (f" — {c.localizare}" if c.localizare else "")
        row[0].text = eticheta
        row[1].text = _b2(pret)
        row[2].text = _pct(bruta)
        # #6 (verificare B): la fel ca la piata — marcam cele N randuri mediate (indici_mediati).
        marcaj = "DA" if i in lr.indici_mediati else ""
        if i == lr.index_selectat:
            marcaj = (marcaj + " *").strip()
        row[3].text = marcaj
    doc.add_paragraph(
        f"Valoarea terenului = {_b2(lr.pret_mp_ales)} EUR/mp (media celor mai similare "
        f"{len(lr.indici_mediati)} comparabile; * = cel mai similar) x {ctx.land.suprafata} mp = "
        f"{_fmt(lr.valoare_teren)} EUR."
    )


def _adauga_tabel_cost(doc: DocxDocument, ctx: ReportContext) -> None:
    if not ctx.building.elements:
        return
    doc.add_paragraph("Tabelul costurilor segregate (abordarea prin cost):")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Element"
    hdr[1].text = "Cod"
    hdr[2].text = "Cantitate"
    hdr[3].text = "Cost unitar"
    hdr[4].text = "Cost de nou"
    for el in ctx.building.elements:
        row = table.add_row().cells
        row[0].text = el.element
        row[1].text = el.cod
        row[2].text = _b2(el.cantitate)
        row[3].text = _b2(el.cost_unitar)
        row[4].text = _fmt(el.cost_nou())
    if ctx.cost_result is not None:
        c = ctx.cost_result
        doc.add_paragraph(
            f"CIB = {_fmt(c.cib)}; Vcp = {_b2(c.vcp)} ani; depreciere fizica = "
            f"{_pct(c.depreciere_fizica)}; CIN = {_fmt(c.cin)}; "
            f"valoare prin cost (CIN + teren) = {_fmt(c.valoare_cost)} {ctx.meta.moneda}."
        )


# --------------------------------------------------------------------------- #
# Back matter (alocare, risc GEV 520, anexe, semnatura)
# --------------------------------------------------------------------------- #
def _adauga_alocare(doc: DocxDocument, ctx: ReportContext, adnotari: bool = False) -> None:
    if ctx.alocare_constructii is None:
        return
    doc.add_heading("ALOCAREA VALORII", level=1)
    _nota(doc, "alocare", adnotari)
    vt = _valoare_teren(ctx)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Componenta"
    table.rows[0].cells[1].text = f"Valoare ({ctx.meta.moneda})"
    randuri = [
        ("Valoarea proprietatii (estimata)", ctx.reconciled.valoare_finala),
        ("din care: Valoarea terenului", vt),
        ("din care: Valoarea constructiilor (alocata)", ctx.alocare_constructii),
    ]
    for nume, val in randuri:
        r = table.add_row().cells
        r[0].text = nume
        r[1].text = _fmt(val) if val is not None else "-"
    doc.add_paragraph(
        "Valoarea constructiilor a fost obtinuta prin alocare: valoarea proprietatii minus "
        "valoarea terenului estimata prin comparatie directa."
    )
    # Verificare de consistenta cost <-> piata (GEV 520): depreciere implicita neexplicata.
    cr = ctx.cost_result
    if cr is not None and cr.cin and cr.cin > 0 and ctx.alocare_constructii is not None:
        deprec = (cr.cin - ctx.alocare_constructii) / cr.cin
        if abs(deprec) > Decimal("0.20"):
            p = doc.add_paragraph()
            p.add_run(
                f"VERIFICARE DE CONSISTENTA (GEV 520): valoarea constructiilor alocata "
                f"({_fmt(ctx.alocare_constructii)} {ctx.meta.moneda}) difera cu {_pct(deprec)} fata de "
                f"costul de inlocuire net ({_fmt(cr.cin)} {ctx.meta.moneda}). O diferenta de aceasta "
                "amplitudine trebuie justificata explicit (depreciere fizica/functionala/externa, "
                "supraoferta sau localizare inferioara) - altfel poate fi semnalata la verificarea bancara."
            ).bold = True


def _adauga_clauza_subasigurare(doc: DocxDocument, ctx: ReportContext, adnotari: bool = False) -> None:
    """Clauza de subasigurare (SEV 450, par. 4 — regula proportionalitatii). Specifica evaluarii pentru
    asigurare: avertizeaza ca o suma asigurata sub costul de reconstructie reduce proportional despagubirea."""
    doc.add_heading("CLAUZA DE SUBASIGURARE (SEV 450)", level=1)
    doc.add_paragraph(
        "Suma asigurata ar trebui sa corespunda valorii de asigurare (costul de reconstructie / de inlocuire "
        "BRUT) estimate in prezentul raport. Daca, la data producerii unei daune, suma asigurata este mai "
        "mica decat costul de reconstructie, despagubirea se reduce PROPORTIONAL cu raportul dintre suma "
        "asigurata si costul de reconstructie (regula proportionalitatii / a subasigurarii), conform "
        "conditiilor politei de asigurare."
    )
    doc.add_paragraph(
        "Se recomanda actualizarea periodica a sumei asigurate, pentru a reflecta evolutia costurilor de "
        "constructie si a evita subasigurarea."
    )


def _adauga_risc_garantie(doc: DocxDocument, ctx: ReportContext, adnotari: bool = False) -> None:
    doc.add_heading("RISCUL ASOCIAT GARANTIEI (GEV 520)", level=1)
    _nota(doc, "gev520", adnotari)
    txt = _narativ(ctx, "Riscul asociat garantiei (GEV 520)")
    doc.add_paragraph(
        txt or
        "Evaluarea pentru garantarea imprumutului respecta cerintele Ghidului GEV 520. Au fost "
        "analizati factorii relevanti pentru estimarea de catre creditor a performantei garantiei "
        "pe perioada creditului, precum si gradul de adecvare al proprietatii ca garantie."
    )
    # GEV 520, Anexa A5: factorii obligatorii de comentat
    doc.add_paragraph("Factori relevanti pentru performanta garantiei (GEV 520, A5):")
    for f in (
        "activitatea curenta si tendintele pietei relevante;",
        "cererea anterioara, curenta si viitoare pentru tipul de bun si pentru localizare;",
        "cererea potentiala sau probabila pentru alte utilizari, la data evaluarii;",
        "impactul evenimentelor previzibile la data evaluarii asupra valorii viitoare a garantiei "
        "pe perioada creditului.",
    ):
        doc.add_paragraph(f, style="List Bullet")
    doc.add_paragraph(
        "Lichiditate si vandabilitate: se apreciaza lichiditatea pietei locale, gradul de adecvare "
        "al proprietatii ca garantie si perioada de comercializare estimata."
    )
    # Inregistrare BIG — conditionat de utilizatorul desemnat (GEV 520 ed. 2025, par. 77-78): raportul
    # de garantare la REESALONAREA datoriilor (utilizator desemnat ANAF) NU se inregistreaza in BIG.
    if ctx.meta.utilizator_desemnat == "ANAF":
        doc.add_paragraph(
            "Inregistrare BIG: raportul are utilizatorul desemnat ANAF (garantare in procesul de "
            "reesalonare a datoriilor, GEV 520 par. 77-78); conform Ghidului, acest tip de raport NU se "
            "inregistreaza in Baza Imobiliara de Garantii (BIG) a ANEVAR."
        )
    else:
        doc.add_paragraph(
            "Inregistrare BIG: raportul, avand utilizarea desemnata de garantare a imprumutului, se "
            "inregistreaza in Baza Imobiliara de Garantii (BIG), conform reglementarilor ANEVAR."
        )
    # B4 — Valoarea de lichidare / vanzare fortata (ceruta frecvent la garantare).
    vp = ctx.reconciled.valoare_finala
    factor = Decimal("0.85")
    p = doc.add_paragraph()
    p.add_run("Valoarea de lichidare (vanzare fortata): ").bold = True
    p.add_run(
        f"valoarea de piata ({_fmt(vp)} {ctx.meta.moneda}) ajustata cu un factor care reflecta "
        f"vanzarea intr-un termen limitat. Estimare orientativa la factor {factor}: "
        f"{_fmt(vp * factor)} {ctx.meta.moneda}. "
    )
    p.add_run(
        "ATENTIE: factorul final se stabileste de evaluator in functie de lichiditatea pietei si "
        "perioada de expunere (tipic 0,80–0,90); valoarea de mai sus este doar orientativa."
    ).italic = True
    # B5 — Anexa de certificare a conformitatii (GEV 520). De confirmat fata de Anexa 1 a ghidului.
    doc.add_heading("CERTIFICAREA CONFORMITATII RAPORTULUI (GEV 520)", level=1)
    doc.add_paragraph(
        "Checklist de conformitate — de verificat si bifat de evaluator inainte de transmitere "
        "(de aliniat la Anexa 1 a GEV 520 in vigoare):"
    )
    # Punctul de checklist BIG — conditionat de utilizatorul desemnat (GEV 520 par. 77-78).
    big_punct = (
        "Raportul are utilizator desemnat ANAF -> NU se inregistreaza in BIG (GEV 520 par. 77-78)."
        if ctx.meta.utilizator_desemnat == "ANAF"
        else "Raportul se inregistreaza in BIG."
    )
    for punct in (
        "Valoarea estimata este valoarea de piata (SEV 102 / IVS 104), exprimata fara TVA.",
        "Data evaluarii si data raportului sunt precizate.",
        "Scopul evaluarii (garantarea imprumutului) este declarat in termenii de referinta.",
        "Proprietatea a fost identificata (cadastral / CF) si inspectata.",
        "Au fost utilizate minim 3 comparabile relevante.",
        "Ajustarile sunt justificate; ajustarile totale se incadreaza in limite rezonabile.",
        "Costul de inlocuire (unde se aplica) exclude profitul dezvoltatorului.",
        "Ipoteza de utilizare continua si ipotezele/conditiile limitative sunt precizate.",
        "A fost analizata cea mai buna utilizare (CMBU).",
        "A fost analizat riscul de garantie (lichiditate, vandabilitate, perioada de expunere).",
        "Valoarea de lichidare / vanzare fortata a fost estimata.",
        "Comparabilele provenite din oferte au fost ajustate la nivel de tranzactie.",
        "Sursele de date sunt mentionate si verificabile.",
        big_punct,
        "Evaluatorul este autorizat ANEVAR si detine asigurare de raspundere profesionala.",
        "Declaratia de conformitate si de independenta este semnata.",
    ):
        doc.add_paragraph(f"☐ {punct}", style="List Bullet")


def _decode_foto(data_url: str) -> BytesIO | None:
    """Extrage bytes dintr-un data-URL base64 (sau base64 simplu). None daca e invalid."""
    if not data_url:
        return None
    # `and "," in ...`: un data-URL fara virgula (ex. "data:") NU se desparte -> ramane intreg si
    # cade pe b64decode care il respinge (None), in loc de IndexError pe [1] -> 500 (RUNDA 9).
    payload = data_url.split(",", 1)[1] if data_url.startswith("data:") and "," in data_url else data_url
    try:
        raw = base64.b64decode(payload, validate=True)
    except (binascii.Error, ValueError):
        return None
    if not raw:
        return None
    return BytesIO(raw)


def _adauga_anexe(doc: DocxDocument, ctx: ReportContext, adnotari: bool = False) -> None:
    doc.add_heading("ANEXE", level=1)
    _nota(doc, "anexe", adnotari)
    doc.add_paragraph("Anexa 1 — Comparabile utilizate (surse):")
    surse = [c.sursa for c in ctx.comparables if c.sursa and c.sursa != "manual"]
    if surse:
        for u in surse:
            doc.add_paragraph(u, style="List Bullet")
    else:
        doc.add_paragraph("Comparabile introduse manual.", style="List Bullet")

    doc.add_paragraph("Anexa 2 — Planse fotografice ale proprietatii:")
    inserate = 0
    for i, data_url in enumerate(ctx.photos, 1):
        stream = _decode_foto(data_url)
        if stream is None:
            log.warning("Anexa 2: fotografia %d nu a putut fi decodată (date invalide) — sărită.", i)
            continue
        try:
            doc.add_picture(stream, width=Inches(5.5))
            inserate += 1
        except Exception as e:   # fisier imagine corupt / format neacceptat -> sarim, dar lasam urma
            log.warning("Anexa 2: fotografia %d nu a putut fi inserată (corupt/format neacceptat): %s", i, e)
            continue
    if inserate == 0:
        doc.add_paragraph("[de atasat].")

    doc.add_paragraph("Anexa 3 — Documente cadastrale, extras CF si acte juridice:")
    doc_inserate = 0
    for i, data_url in enumerate(ctx.documente, 1):
        stream = _decode_foto(data_url)
        if stream is None:
            log.warning("Anexa 3: documentul %d nu a putut fi decodat (date invalide) — sărit.", i)
            continue
        try:
            doc.add_picture(stream, width=Inches(5.5))
            doc_inserate += 1
        except Exception as e:
            log.warning("Anexa 3: documentul %d nu a putut fi inserat (corupt/format neacceptat): %s", i, e)
            continue
    if doc_inserate == 0:
        doc.add_paragraph("[de atasat].")


def _adauga_semnatura(doc: DocxDocument, ctx: ReportContext, adnotari: bool = False) -> None:
    meta = ctx.meta
    _nota(doc, "semnatura", adnotari)
    doc.add_paragraph("")
    doc.add_paragraph("Intocmit,")
    p = doc.add_paragraph()
    p.add_run(meta.evaluator_nume).bold = True
    doc.add_paragraph(f"Evaluator autorizat ANEVAR, legitimatia {meta.evaluator_legitimatie}")
    doc.add_paragraph(f"Data: {meta.data_raportului}")
    doc.add_paragraph("Semnatura / stampila: ______________________")


# --------------------------------------------------------------------------- #
# Asamblare
# --------------------------------------------------------------------------- #
def genereaza_raport(
    ctx: ReportContext, output_path: Path | str, adnotari: bool = False
) -> Path:
    """Construieste si salveaza raportul .docx. Returneaza calea fisierului.

    Cu `adnotari=True`, insereaza note de provenienta (mod demo/review) sub fiecare sectiune:
    ce e calculat real, extras automat, generat de AI, exemplu sau placeholder.
    """
    doc = Document()
    meta = ctx.meta

    # Disclaimer al aplicatiei catre evaluator — in fruntea ORICARUI raport generat (cerinta Adi).
    _disclaimer_aplicatie(doc)

    # --- Shell GBF (front matter) ---
    _coperta(doc, ctx, adnotari)
    _scrisoare_transmitere(doc, ctx, adnotari)
    _declaratie_conformitate(doc, ctx, adnotari)
    _termeni_referinta(doc, ctx, adnotari)

    # --- Cele 7 capitole de analiza (raportare conform SEV 106 Documentare si raportare) ---
    doc.add_heading("1. SINTEZA EVALUARII SI CERTIFICARE", level=1)
    _nota(doc, "cap1", adnotari)
    doc.add_paragraph(f"Client: {meta.client_nume} ({meta.client_tip}).")
    doc.add_paragraph(
        f"Proprietatea: {meta.adresa}; nr. cadastral {meta.numar_cadastral}; "
        f"{meta.carte_funciara}."
    )
    doc.add_paragraph(f"Scopul evaluarii: {meta.scop}.")
    doc.add_paragraph(f"Tipul valorii: {_tip_valoare_txt(meta.tip_valoare)}.")
    doc.add_paragraph(
        f"Data evaluarii: {meta.data_evaluarii}; data raportului: {meta.data_raportului}."
    )
    doc.add_paragraph(
        f"VALOAREA ESTIMATA: {_fmt(ctx.reconciled.valoare_finala)} {meta.moneda} "
        f"(metoda selectata: {ctx.reconciled.metoda_selectata}). {_fara_tva(ctx)}"
    )

    doc.add_heading("2. IPOTEZE GENERALE SI SPECIALE", level=1)
    _nota(doc, "cap2", adnotari)
    doc.add_paragraph(
        _narativ(ctx, "Ipoteze generale si speciale")
        or "Ipoteze limitative standard privind structura de rezistenta si solul."
    )

    doc.add_heading("3. PREZENTAREA DATELOR DE PIATA", level=1)
    _nota(doc, "cap3", adnotari)
    doc.add_paragraph(
        _narativ(ctx, "Prezentarea datelor de piata")
        or "Analiza pietei locale [de completat]."
    )

    doc.add_heading("4. DESCRIEREA JURIDICA SI FIZICA A PROPRIETATII", level=1)
    _nota(doc, "cap4", adnotari)
    doc.add_paragraph(f"Teren: {ctx.land.suprafata} mp, categorie {ctx.land.categorie}.")
    if ctx.land.categorie_folosinta is not None:
        clasa = f", clasa de calitate {ctx.land.clasa_calitate}" if ctx.land.clasa_calitate is not None else ""
        doc.add_paragraph(f"Teren agricol: categorie de folosinta {ctx.land.categorie_folosinta}{clasa}.")
    if ctx.land.acces:                                  # GEV 630 §28 — calea de acces (afecteaza valoarea)
        doc.add_paragraph(f"Acces: {ctx.land.acces}.")
    if ctx.land.utilitati:                              # GEV 630 §28 — utilitati in descriere
        doc.add_paragraph(f"Utilitati: {', '.join(ctx.land.utilitati)}.")
    if ctx.land.restrictii_urbanism:                    # GEV 630 §16 — regim/restrictii urbanism (POT/CUT etc.)
        doc.add_paragraph(f"Regim urbanistic / restrictii (certificat urbanism): {ctx.land.restrictii_urbanism}.")
    doc.add_paragraph(
        f"Constructie: Au {ctx.building.au} mp, Acd {ctx.building.acd} mp, "
        f"an referinta {ctx.building.an_referinta}."
    )
    if ctx.building.etaj is not None or ctx.building.an_bloc is not None:
        parti = []
        if ctx.building.etaj is not None:
            niv = f"/{ctx.building.nr_niveluri_bloc}" if ctx.building.nr_niveluri_bloc else ""
            parti.append(f"etaj {ctx.building.etaj}{niv}")
        if ctx.building.an_bloc is not None:
            parti.append(f"an bloc {ctx.building.an_bloc}")
        if ctx.building.cota_teren_indiviza is not None:
            parti.append(f"cota teren indiviza {ctx.building.cota_teren_indiviza} mp")
        doc.add_paragraph("Apartament: " + ", ".join(parti) + ".")
    if ctx.building.inaltime_libera is not None:
        doc.add_paragraph(f"Spatiu industrial: inaltime libera {ctx.building.inaltime_libera} m.")
    insp = []
    if meta.data_inspectiei:
        insp.append(f"data inspecției {meta.data_inspectiei}")
    if meta.inspectie_amploare:
        insp.append(f"amploare {meta.inspectie_amploare}")
    if meta.inspectie_insotitor:
        insp.append(f"însoțit de {meta.inspectie_insotitor}")
    doc.add_paragraph("Inspectia proprietatii: " + "; ".join(insp) + "." if insp
                      else "Inspectia proprietatii: amploarea si insotitorul nu au fost declarate (de completat).")
    if meta.inspectie_observatii:
        doc.add_paragraph(f"Observatii / neconcordante scriptic-faptic la inspectie: {meta.inspectie_observatii}.")
    descriere = _narativ(ctx, "Descrierea juridica si fizica a proprietatii")
    if descriere:
        doc.add_paragraph(descriere)

    doc.add_heading("5. ANALIZA CELEI MAI BUNE UTILIZARI (CMBU)", level=1)
    _nota(doc, "cap5", adnotari)
    doc.add_paragraph(
        _narativ(ctx, "Analiza celei mai bune utilizari (CMBU)")
        or "Analiza CMBU [de completat]."
    )

    doc.add_heading("6. APLICAREA METODELOR DE CALCUL", level=1)
    _nota(doc, "cap6", adnotari)
    doc.add_paragraph(
        'Abordările și metodele aplicate respectă SEV 103 „Abordări în evaluare” și '
        'SEV 105 „Modele de evaluare” (echivalente IVS 103/IVS 105). Selecția abordării '
        'este justificată în reconcilierea de la capitolul 7.'
    )
    _adauga_grila_comparatie(doc, ctx)
    _adauga_grila_teren(doc, ctx)
    _adauga_tabel_cost(doc, ctx)
    if ctx.venit_result is not None:
        doc.add_heading("Abordarea prin venit (capitalizare directă)", level=2)
        dv, vr = ctx.date_venit, ctx.venit_result
        if dv is not None:
            doc.add_paragraph(
                f"Venit brut potențial: {dv.venit_brut_potential} lei/an; "
                f"neocupare {dv.grad_neocupare}; cheltuieli {dv.cheltuieli_exploatare} lei/an; "
                f"rată de capitalizare {dv.rata_capitalizare}."
            )
        doc.add_paragraph(
            f"Venit net din exploatare (NOI): {vr.noi} lei. "
            f"Valoare = NOI / rată = {vr.valoare} lei."
        )
    if ctx.dcf_valoare is not None:
        doc.add_heading("Abordarea prin venit (DCF)", level=2)
        doc.add_paragraph(f"Valoare prin actualizarea fluxurilor de numerar: {ctx.dcf_valoare} lei.")
    justificare = _narativ(ctx, "Justificarea ajustarilor aplicate")
    if justificare:
        doc.add_paragraph(justificare)

    doc.add_heading("7. RECONCILIEREA REZULTATELOR SI CONCLUZIA VALORII", level=1)
    _nota(doc, "cap7", adnotari)
    doc.add_paragraph(
        _narativ(ctx, "Reconcilierea rezultatelor si concluzia valorii")
        or "Reconcilierea metodelor [de completat]."
    )
    doc.add_paragraph(
        f"Valoarea finala: {_fmt(ctx.reconciled.valoare_finala)} {meta.moneda}"
        f"{_echiv_lei(ctx)}. {_fara_tva(ctx)}"
    )
    if ctx.reconciled.nota:   # transparenta reconciliere (ex. abordare calculata dar neponderata)
        doc.add_paragraph(f"Notă privind reconcilierea: {ctx.reconciled.nota}")

    # --- Shell GBF (back matter) ---
    _adauga_alocare(doc, ctx, adnotari)
    if ctx.profil.ghid == "GEV_520":   # sectiunea de risc e specifica garantarii imprumutului
        _adauga_risc_garantie(doc, ctx, adnotari)
    if ctx.profil.ghid == "SEV_450":   # clauza de subasigurare e specifica evaluarii pentru asigurare
        _adauga_clauza_subasigurare(doc, ctx, adnotari)
    _adauga_anexe(doc, ctx, adnotari)
    _adauga_semnatura(doc, ctx, adnotari)

    output_path = Path(output_path)
    doc.save(str(output_path))
    # Log operațional (fără PII: doar nume fișier=uid, metodă, nr. anexe cerute) — diagnoză pe teren.
    log.info("Raport generat: %s (metoda=%s, foto cerute=%d, scanuri cerute=%d)",
             output_path.name, ctx.reconciled.metoda_selectata,
             len(ctx.photos), len(ctx.documente))
    return output_path
