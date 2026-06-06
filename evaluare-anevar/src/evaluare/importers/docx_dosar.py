"""Import dosar dintr-un raport `.docx` existent -> câmpuri wizard (pre-completare).

Strategie pe două straturi (robustă la rapoarte cu format liber):
  1. **Numele fișierului** dă identitatea sigură: `<id> <nume client> <tip> <localitate>`
     (ex. `21766 Bololoi Daniela-Doina locuinta Busteni.docx`).
  2. **Textul** dă extra (best-effort, prin regex): beneficiar/bancă, scop, data inspecției,
     județul (din localitate). Câmpurile negăsite rămân goale — fără ghicire tăcută.

Folosit la „Importă dosarul tău" și ca temelie pentru import „dosar asemănător" (feature D).
Vezi docs/specs/1-ui-output-first.md și docs/features-majore-pending.md.
"""
from __future__ import annotations

import re
from pathlib import Path

# Tip din numele fișierului -> tipul intern (Pas 2). „locuinta"/„casa" -> casa.
_TIP = {
    "casa": "casa", "locuinta": "casa", "locuință": "casa", "vila": "casa",
    "apartament": "apartament", "ap": "apartament", "garsoniera": "apartament",
    "teren": "agricol", "agricol": "agricol", "hala": "industrial", "industrial": "industrial",
    "spatiu": "special", "comercial": "special",
}
_TIPURI_RE = "|".join(sorted(_TIP, key=len, reverse=True))

# Localitate -> județ (cele uzuale din exemplele de portofoliu; extensibil).
_JUDET = {
    "busteni": "Prahova", "breaza": "Prahova", "maneciu": "Prahova", "ploiesti": "Prahova",
    "sinaia": "Prahova", "campina": "Prahova", "azuga": "Prahova", "comarnic": "Prahova",
    "brasov": "Brașov", "predeal": "Brașov", "rasnov": "Brașov", "zarnesti": "Brașov",
    "bucuresti": "București", "constanta": "Constanța", "cluj": "Cluj", "cluj-napoca": "Cluj",
}


def parse_nume_fisier(nume: str) -> dict:
    """Extrage identitatea din numele fișierului: id_client, nume_client, tip, localitate."""
    baza = Path(nume).stem.strip()
    out: dict[str, str] = {}
    m = re.match(r"^\s*(\d{3,})\s+(.*)$", baza)          # id numeric la început
    if m:
        out["id_client"] = m.group(1)
        rest = m.group(2)
    else:
        rest = baza
    # tipul (primul cuvânt-cheie din listă) împarte „nume client" de „localitate"
    mt = re.search(rf"\b({_TIPURI_RE})\b", rest, re.IGNORECASE)
    if mt:
        out["tip_proprietate"] = _TIP[mt.group(1).lower()]
        nume_client = rest[: mt.start()].strip(" -,")
        localitate = rest[mt.end():].strip(" -,.")
    else:
        nume_client, localitate = rest.strip(), ""
    if nume_client:
        out["nume_client"] = nume_client
    # curăță localitatea: păstrează primul token „cuvânt" (taie „str.Fantanii ,Brasov" -> Brasov)
    if localitate:
        loc = re.split(r"[,/]", localitate)
        cand = loc[-1].strip()                          # ce e după virgulă e de obicei orașul
        cand = re.sub(r"^\s*str\.?\s*\w+\s*", "", cand, flags=re.IGNORECASE).strip()
        localitate = cand or loc[0].strip()
        localitate = re.sub(r"\s*-\s*", "-", localitate).strip(" -")   # „Maneciu -Pamanteni" -> „Maneciu-Pamanteni"
        out["localitate"] = localitate.title()
        prim = re.split(r"[\s-]", localitate.lower(), maxsplit=1)[0]   # „cluj-napoca" -> „cluj"
        jud = _JUDET.get(re.sub(r"[^a-z]", "", prim)) or _JUDET.get(re.sub(r"[^a-z-]", "", localitate.lower()))
        if jud:
            out["judet"] = jud
    return out


def _text(path: Path) -> str:
    import docx
    doc = docx.Document(str(path))
    parti = [p.text for p in doc.paragraphs]
    for t in doc.tables:                                # și textul din tabele
        for row in t.rows:
            parti += [c.text for c in row.cells]
    return "\n".join(parti)


def extrage_din_text(text: str) -> dict:
    """Best-effort din corpul raportului: beneficiar/bancă, scop, data inspecției."""
    out: dict[str, str] = {}
    mb = re.search(r"\b([A-ZĂÂÎȘȚ][A-Za-zĂÂÎȘȚăâîșț]+\s+BANK(?:\s+[A-Z]{2,})?)\b", text)
    if mb:
        out["beneficiar"] = re.sub(r"\s+", " ", mb.group(1)).strip().upper()
    if re.search(r"garantarea\s+(împrumutului|creditului)", text, re.IGNORECASE):
        out["scop"] = "garantare"
    elif re.search(r"raportare\s+financiar", text, re.IGNORECASE):
        out["scop"] = "raportare_financiara"
    md = re.search(r"[Ii]nspec[țt]i[ae][^\d]{0,40}?(\d{2}[.\-/]\d{2}[.\-/]\d{4})", text)
    if md:
        zi, luna, an = re.split(r"[.\-/]", md.group(1))
        out["data_inspectiei"] = f"{an}-{luna}-{zi}"
    return out


def extrage_din_docx(path: str | Path) -> dict:
    """Câmpurile wizard pre-completate dintr-un raport `.docx` (filename + text)."""
    import contextlib
    p = Path(path)
    wizard = parse_nume_fisier(p.name)
    wizard.setdefault("scop", "garantare")
    with contextlib.suppress(Exception):                # docx ilizibil -> rămânem pe filename
        wizard.update({k: v for k, v in extrage_din_text(_text(p)).items() if v})
    return wizard
