"""Generatoare de documente AML (.docx) — reutilizeaza python-docx.

Documente: norme interne (7 capitole, Norme art. 8(1) a-g), evaluarea proprie de risc,
decizie de desemnare (DOAR societate), fisa KYC, draft RTN, draft RTS (cu avertisment
tipping-off). Continutul e pre-completat; evaluatorul verifica si semneaza/transmite.
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional, Union

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt

from evaluare.aml.incadrare import necesita_persoana_desemnata
from evaluare.aml.models import (
    ClientPF, ClientPJ, EvaluareRisc, PersoanaFizica, StatutPEP,
)
from evaluare.aml.raportare import RaportRTN, RaportRTS

Client = Union[ClientPF, ClientPJ]


def _antet(doc: DocxDocument, antet: Optional[dict]) -> None:
    """Antet uzual: entitate, data, versiune, clauza de aprobare/revizuire."""
    antet = antet or {}
    entitate = antet.get("entitate", "______________________ (entitatea raportoare)")
    data = antet.get("data", "____________")
    versiune = antet.get("versiune", "1.0")
    p = doc.add_paragraph()
    p.add_run(f"Entitate raportoare: {entitate}").bold = True
    doc.add_paragraph(f"Data: {data}    Versiune: {versiune}")
    doc.add_paragraph(
        "Aprobat la nivelul conducerii de rang superior. Document elaborat în temeiul Legii "
        "nr. 129/2019 și al Normelor aprobate prin Ordinul ONPCSB nr. 37/2021. Se revizuiește "
        "ori de câte ori intervin modificări legislative sau organizatorice."
    )


def _kv(doc: DocxDocument, eticheta: str, valoare) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{eticheta}: ").bold = True
    p.add_run("" if valoare is None else str(valoare))


def _pf_linie(p: PersoanaFizica) -> str:
    buc = [f"{p.nume} {p.prenume}".strip()]
    if p.cnp:
        buc.append(f"CNP {p.cnp}")
    act = " ".join(x for x in [p.tip_act, p.serie_act, p.nr_act] if x)
    if act:
        buc.append(f"act: {act}")
    if p.cetatenie:
        buc.append(p.cetatenie)
    if p.domiciliu:
        buc.append(p.domiciliu)
    return "; ".join(buc)


# --------------------------------------------------------------------------- #
# 1. Norme interne (7 capitole, Norme art. 8(1) a-g)
# --------------------------------------------------------------------------- #
_CAPITOLE_NORME = [
    ("1. Raportare și păstrarea evidențelor",
     "Măsuri de raportare către autoritățile statului (RTS, RTN, raportări la solicitare), "
     "inclusiv semnalarea în nume propriu a încălcărilor, și de păstrare a evidențelor și "
     "documentelor conform cerințelor legale (retenție 5 ani). — Norme art. 8(1)(a); Legea art. 6-9, 21."),
    ("2. Cunoașterea clientelei (KYC)",
     "Măsuri de identificare și verificare a clientului și a beneficiarului real, momentul aplicării, "
     "documentele acceptate, măsuri standard/simplificate/suplimentare. — Norme art. 8(1)(b); Legea art. 11-17."),
    ("3. Administrarea riscurilor",
     "Identificarea, evaluarea, gestionarea și diminuarea riscurilor SB/FT; criteriile și factorii de "
     "risc; scenariile și intervalele pentru identificarea tranzacțiilor legate. — Norme art. 8(1)(c)."),
    ("4. Control intern, comunicare și management de conformitate",
     "Mecanisme de control intern, fluxuri de comunicare și responsabilități de conformitate. "
     "— Norme art. 8(1)(d)."),
    ("5. Protecția personalului",
     "Măsuri de protecție a personalului implicat împotriva amenințărilor, acțiunilor ostile sau "
     "discriminatorii. — Norme art. 8(1)(e)."),
    ("6. Avertizare internă (whistleblowing)",
     "Proceduri prin care angajații pot raporta intern încălcări printr-un canal specific, independent "
     "și anonim, proporțional cu natura și dimensiunea entității. — Norme art. 8(1)(f)."),
    ("7. Instruire și evaluarea personalului",
     "Programe de instruire periodică și evaluarea angajaților pentru recunoașterea operațiunilor "
     "asociate SB/FT. — Norme art. 8(1)(g)."),
]


def genereaza_norme_interne(antet: Optional[dict] = None) -> DocxDocument:
    doc = Document()
    doc.add_heading("NORME INTERNE DE PREVENIRE A SPĂLĂRII BANILOR ȘI FINANȚĂRII TERORISMULUI", level=0)
    _antet(doc, antet)
    for titlu, continut in _CAPITOLE_NORME:
        doc.add_heading(titlu, level=1)
        doc.add_paragraph(continut)
    doc.add_paragraph()
    doc.add_paragraph(
        "Notă: conținut-cadru generat automat ca punct de plecare; necesită adaptare la specificul "
        "activității și validare juridică înainte de utilizare."
    )
    return doc


# --------------------------------------------------------------------------- #
# 2. Evaluarea proprie de risc
# --------------------------------------------------------------------------- #
def genereaza_evaluare_risc(evaluare: EvaluareRisc, antet: Optional[dict] = None) -> DocxDocument:
    doc = Document()
    doc.add_heading("EVALUAREA DE RISC A RELAȚIEI DE AFACERI", level=0)
    _antet(doc, antet)
    _kv(doc, "Categorie de risc", evaluare.categorie)
    _kv(doc, "Nivel măsuri de cunoaștere", evaluare.nivel_masuri)
    _kv(doc, "Scor", evaluare.scor)
    _kv(doc, "Data evaluării", evaluare.data)
    _kv(doc, "Data reevaluării", evaluare.data_reevaluare)

    doc.add_heading("Factori de risc (Norme art. 12-13)", level=1)
    tab = doc.add_table(rows=1, cols=3)
    tab.style = "Light Grid Accent 1"
    h = tab.rows[0].cells
    h[0].text, h[1].text, h[2].text = "Factor", "Valoare (1-3)", "Pondere"
    for f in evaluare.factori:
        r = tab.add_row().cells
        r[0].text, r[1].text, r[2].text = f.nume, str(f.valoare), str(f.pondere)

    if evaluare.motive_sporit:
        doc.add_heading("Motive pentru risc sporit (Legea art. 17)", level=1)
        for m in evaluare.motive_sporit:
            doc.add_paragraph(m, style="List Bullet")
    return doc


# --------------------------------------------------------------------------- #
# 3. Decizie de desemnare a persoanei responsabile — DOAR societate
# --------------------------------------------------------------------------- #
def genereaza_decizie_desemnare(
    tip_entitate: str, persoana: PersoanaFizica, antet: Optional[dict] = None,
) -> DocxDocument:
    if not necesita_persoana_desemnata(tip_entitate):
        raise ValueError(
            "PFA / persoană fizică nu are obligația desemnării unei persoane responsabile "
            "(Norme art. 7; Legea art. 23(4)). Decizia nu se generează."
        )
    doc = Document()
    doc.add_heading("DECIZIE DE DESEMNARE A PERSOANEI RESPONSABILE AML/CFT", level=0)
    _antet(doc, antet)
    doc.add_paragraph(
        "În temeiul art. 23 alin. (1)-(3) din Legea nr. 129/2019 și al Normelor ONPCSB nr. 37/2021, "
        "se desemnează persoana responsabilă cu aplicarea prevederilor legale privind prevenirea "
        "spălării banilor și finanțării terorismului, cu următoarele atribuții (HCD 58 Anexa 1):"
    )
    _kv(doc, "Persoana desemnată", _pf_linie(persoana))
    for atrib in [
        "a) elaborarea/actualizarea normelor și procedurilor interne;",
        "b) primirea și analiza rapoartelor interne de tranzacții suspecte;",
        "c) raportarea către ONPCSB (RTS, RTN, raportări la solicitare);",
        "d) relația cu ONPCSB și cu organele de control;",
        "e) instruirea personalului;",
        "f) păstrarea evidențelor și asigurarea confidențialității (tipping-off).",
    ]:
        doc.add_paragraph(atrib, style="List Bullet")
    doc.add_paragraph("\nData: ____________     Semnătura conducerii: ____________________")
    return doc


# --------------------------------------------------------------------------- #
# 4. Fisa KYC
# --------------------------------------------------------------------------- #
def _pep_text(pep: StatutPEP) -> str:
    if not pep.este_pep:
        return "Nu"
    buc = ["Da"]
    if pep.categorie:
        buc.append(pep.categorie)
    if pep.tip:
        buc.append(pep.tip)
    if pep.data_incetare_functie:
        buc.append(f"încetare {pep.data_incetare_functie}")
    return "; ".join(buc)


def genereaza_fisa_kyc(
    client: Client, evaluare: Optional[EvaluareRisc] = None, antet: Optional[dict] = None,
) -> DocxDocument:
    doc = Document()
    doc.add_heading("FIȘĂ DE CUNOAȘTERE A CLIENTELEI (KYC)", level=0)
    _antet(doc, antet)

    if isinstance(client, ClientPF):
        doc.add_heading("Client — persoană fizică", level=1)
        _kv(doc, "Identificare", _pf_linie(client.persoana))
        _kv(doc, "Ocupație", client.persoana.ocupatie)
        _kv(doc, "PEP", _pep_text(client.pep))
    else:
        doc.add_heading("Client — persoană juridică", level=1)
        _kv(doc, "Denumire", client.denumire)
        _kv(doc, "CUI", client.cui)
        _kv(doc, "Sediu", client.sediu)
        _kv(doc, "Reprezentant legal", _pf_linie(client.reprezentant_legal))
        if client.imputernicit:
            _kv(doc, "Împuternicit", _pf_linie(client.imputernicit))
        if client.tip == "PJ_straina":
            _kv(doc, "Traducere legalizată acte", "Da" if client.traducere_legalizata else "NU — necesară (art. 15(2))")
        doc.add_heading("Beneficiari reali (Legea art. 4)", level=2)
        if not client.beneficiari_reali:
            doc.add_paragraph("De completat — identificarea beneficiarului real este obligatorie.")
        for br in client.beneficiari_reali:
            procent = "" if br.procent is None else f" — {br.procent * 100:.0f}%"
            ctrl = f" ({br.tip_control})" if br.tip_control else ""
            doc.add_paragraph(f"{_pf_linie(br)}{procent}{ctrl}; PEP: {_pep_text(br.pep)}",
                              style="List Bullet")

    if evaluare is not None:
        doc.add_heading("Rezultatul evaluării de risc", level=1)
        _kv(doc, "Categorie", evaluare.categorie)
        _kv(doc, "Măsuri", evaluare.nivel_masuri)

    doc.add_paragraph("\nVerificat de: ____________   Data: __________   Semnătura: __________")
    return doc


# --------------------------------------------------------------------------- #
# 5. Draft RTN (numerar)
# --------------------------------------------------------------------------- #
def genereaza_rtn(raport: RaportRTN, antet: Optional[dict] = None) -> DocxDocument:
    doc = Document()
    doc.add_heading("RAPORT PENTRU TRANZACȚII CU NUMERAR (RTN)", level=0)
    _antet(doc, antet)
    _kv(doc, "Sumă (EUR echivalent)", raport.suma_eur)
    _kv(doc, "Data tranzacției", raport.data_tranzactie)
    _kv(doc, "Termen de transmitere (3 zile lucrătoare)", raport.termen)
    _kv(doc, "Descriere", raport.descriere)
    doc.add_paragraph(
        "\nSe transmite exclusiv ONPCSB (rapoarte.onpcsb.ro). Prag legal: echivalentul în lei a "
        "10.000 € (Legea art. 7(1)). Conversie la cursul BNR de la data tranzacției."
    )
    return doc


# --------------------------------------------------------------------------- #
# 6. Draft RTS (suspiciune) — cu avertisment tipping-off
# --------------------------------------------------------------------------- #
def genereaza_rts(raport: RaportRTS, antet: Optional[dict] = None) -> DocxDocument:
    doc = Document()
    doc.add_heading("RAPORT PENTRU TRANZACȚII SUSPECTE (RTS)", level=0)
    # avertisment vizibil sus
    p = doc.add_paragraph()
    run = p.add_run(raport.avertisment_tipping_off)
    run.bold = True
    run.font.size = Pt(11)
    _antet(doc, antet)
    _kv(doc, "Motivul suspiciunii", raport.motiv)
    _kv(doc, "Data înregistrării interne", raport.data_inregistrare)
    _kv(doc, "Suspendare tranzacție până la", raport.suspendare_pana_la)
    if raport.indicatori:
        doc.add_heading("Indicatori de suspiciune (HCD 58 art. 6(10))", level=1)
        for ind in raport.indicatori:
            doc.add_paragraph(ind, style="List Bullet")
    doc.add_paragraph(
        "\nFlux: angajat → persoana responsabilă → ONPCSB (HCD 58 Anexa 2). Se transmite exclusiv "
        "ONPCSB. Document confidențial, stocat SEPARAT de dosarul de evaluare (Legea art. 38)."
    )
    return doc


# --------------------------------------------------------------------------- #
# salvare
# --------------------------------------------------------------------------- #
def salveaza(doc: DocxDocument, cale: Union[str, Path]) -> Path:
    cale = Path(cale)
    cale.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(cale))
    return cale
