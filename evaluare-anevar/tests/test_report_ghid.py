"""Raportul citeaza ghidul (GEV) aplicabil din profil + sectiunea de risc doar la GEV 520."""
from docx import Document

from evaluare.profil import ProfilEvaluare
from evaluare.report.generator import genereaza_raport

from tests.test_report_generator import _ctx


def _text(path) -> str:
    doc = Document(str(path))
    parti = [p.text for p in doc.paragraphs]
    for tab in doc.tables:
        for row in tab.rows:
            parti.extend(c.text for c in row.cells)
    return "\n".join(parti)


def test_default_gev520_neschimbat(tmp_path):
    out = tmp_path / "r.docx"
    genereaza_raport(_ctx(), out)
    t = _text(out)
    assert "GEV 520" in t
    assert "garantarea imprumutului" in t.lower() or "garantarea imprumutului" in t


def test_profil_gev630_citeaza_gev630_fara_risc_garantie(tmp_path):
    ctx = _ctx()
    ctx.profil = ProfilEvaluare(tip_activ="comercial", ghid="GEV_630",
                                abordari_aplicabile=["venit", "comparatie"])
    out = tmp_path / "r630.docx"
    genereaza_raport(ctx, out)
    t = _text(out)
    assert "GEV 630" in t
    # nu mai apare clauza specifica de garantare a imprumutului
    assert "GEV 520 — Evaluarea pentru garantarea" not in t


def test_profil_gev500_citeaza_gev500(tmp_path):
    ctx = _ctx()
    ctx.profil = ProfilEvaluare(tip_activ="casa", scop="impozitare", ghid="GEV_500",
                                abordari_aplicabile=["cost"])
    out = tmp_path / "r500.docx"
    genereaza_raport(ctx, out)
    assert "GEV 500" in _text(out)
