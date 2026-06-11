from decimal import Decimal

from evaluare.models.meta import EvaluationMeta
from evaluare.models.narrative import NarrativeSection
from evaluare.valoare_prudenta import ParametriValoarePrudenta


def _meta_minim(**extra) -> EvaluationMeta:
    """EvaluationMeta cu setul minim de campuri obligatorii (+ orice override)."""
    return EvaluationMeta(
        client_nume="Ion Popescu",
        adresa="Str. Exemplu nr. 1, Bucuresti",
        numar_cadastral="123456",
        carte_funciara="CF123456",
        evaluator_nume="Maria Ionescu",
        evaluator_legitimatie="19567",
        data_evaluarii="2026-01-16",
        data_raportului="2026-01-16",
        **extra,
    )


def test_evaluation_meta_required_and_defaults():
    m = EvaluationMeta(
        client_nume="Ion Popescu",
        adresa="Str. Exemplu nr. 1, Bucuresti",
        numar_cadastral="123456",
        carte_funciara="CF123456",
        evaluator_nume="Maria Ionescu",
        evaluator_legitimatie="19567",
        data_evaluarii="2026-01-16",
        data_raportului="2026-01-16",
    )
    assert m.client_tip == "Persoana fizica"          # default
    assert m.scop == "Garantarea creditului ipotecar"  # default
    assert m.moneda == "LEI"                            # default
    assert m.client_nume == "Ion Popescu"


def test_evaluation_meta_big_esg_fields_defaults():
    """Campurile BIG/ESG noi sunt optionale -> constructiile existente raman valide.

    Acelasi set minim de campuri ca testul de baza; verificam doar valorile implicite
    ale campurilor adaugate (gap S-4 cod_postal, S-5 riscuri_fizice, G7 CPE).
    """
    m = EvaluationMeta(
        client_nume="Ion Popescu",
        adresa="Str. Exemplu nr. 1, Bucuresti",
        numar_cadastral="123456",
        carte_funciara="CF123456",
        evaluator_nume="Maria Ionescu",
        evaluator_legitimatie="19567",
        data_evaluarii="2026-01-16",
        data_raportului="2026-01-16",
    )
    assert m.cod_postal is None                # default
    assert m.riscuri_fizice == []              # default factory (lista goala)
    assert m.certificat_energetic is None      # default
    # lista implicita e o instanta proprie, nu un default mutabil partajat
    m.riscuri_fizice.append("seismic")
    other = EvaluationMeta(
        client_nume="Alt Client",
        adresa="Str. Alta nr. 2",
        numar_cadastral="654321",
        carte_funciara="CF654321",
        evaluator_nume="Maria Ionescu",
        evaluator_legitimatie="19567",
        data_evaluarii="2026-01-16",
        data_raportului="2026-01-16",
    )
    assert other.riscuri_fizice == []


def test_evaluation_meta_big_esg_fields_set():
    """Campurile BIG/ESG pot fi populate (cod postal, riscuri fizice, CPE)."""
    m = EvaluationMeta(
        client_nume="Ion Popescu",
        adresa="Str. Exemplu nr. 1, Bucuresti",
        numar_cadastral="123456",
        carte_funciara="CF123456",
        evaluator_nume="Maria Ionescu",
        evaluator_legitimatie="19567",
        data_evaluarii="2026-01-16",
        data_raportului="2026-01-16",
        cod_postal="010101",
        riscuri_fizice=["inundabilitate", "seismic"],
        certificat_energetic="Clasa B",
    )
    assert m.cod_postal == "010101"
    assert m.riscuri_fizice == ["inundabilitate", "seismic"]
    assert m.certificat_energetic == "Clasa B"


# ── Valoarea prudenta (CRR III) — campul optional valoare_prudenta_params ─────
def test_valoare_prudenta_params_optional_default_none():
    """Backward-compatible: campul exista, e optional si implicit None.

    None -> sectiunea de valoare prudenta din raport ramane omisa (ca inainte de cablare)."""
    m = _meta_minim()
    assert m.valoare_prudenta_params is None


def test_valoare_prudenta_params_accepta_instanta():
    """Populat cu o instanta `ParametriValoarePrudenta` -> pastrata ca atare (validata)."""
    p = ParametriValoarePrudenta(discount_crestere_pct=Decimal("10"))
    m = _meta_minim(valoare_prudenta_params=p)
    assert isinstance(m.valoare_prudenta_params, ParametriValoarePrudenta)
    assert m.valoare_prudenta_params.discount_crestere_pct == Decimal("10")


def test_valoare_prudenta_params_accepta_dict_si_il_valideaza():
    """Populat cu un dict (ex. din JSON-ul lucrarii) -> coercitionat la model validat."""
    m = _meta_minim(valoare_prudenta_params={"haircut_sustenabilitate_pct": Decimal("5")})
    assert isinstance(m.valoare_prudenta_params, ParametriValoarePrudenta)
    assert m.valoare_prudenta_params.haircut_sustenabilitate_pct == Decimal("5")


def test_raport_garantare_cu_params_randeaza_sectiunea_valoare_prudenta(tmp_path):
    """End-to-end: meta de baza (fara subtip) cu params populat -> raportul de garantare (GEV 520)
    capata sectiunea de valoare prudenta. Inainte de adaugarea campului, sectiunea nu se activa
    niciodata din meta reala (doar dintr-un subtip ad-hoc in teste)."""
    from docx import Document

    from evaluare.models.property import (
        BuildingData,
        CostElement,
        DepreciationPoint,
        LandData,
    )
    from evaluare.models.report_context import ReportContext
    from evaluare.models.results import CostResult, ReconciledResult
    from evaluare.report.generator import genereaza_raport

    meta = _meta_minim(
        valoare_prudenta_params=ParametriValoarePrudenta(discount_crestere_pct=Decimal("10"))
    )
    ctx = ReportContext(
        meta=meta,
        land=LandData(suprafata=Decimal("500"), categorie="intravilan"),
        building=BuildingData(
            au=Decimal("100"), acd=Decimal("120"), an_referinta=2025,
            elements=[CostElement(element="Structura", cod="X", um="mp",
                                  cantitate=Decimal("100"), cost_unitar=Decimal("2000"),
                                  an_pif=2015)],
            depreciation_points=[DepreciationPoint(varsta=10, depreciere=Decimal("0.1"))],
        ),
        cost_result=CostResult(cib=Decimal("200000"), vcp=Decimal("10"),
                               depreciere_fizica=Decimal("0.10"), cin=Decimal("180000"),
                               valoare_cost=Decimal("280000")),
        reconciled=ReconciledResult(valoare_finala=Decimal("280000"), metoda_selectata="cost"),
    )
    out = tmp_path / "garantare_prudenta.docx"
    genereaza_raport(ctx, out)
    doc = Document(str(out))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts)
    assert "VALOAREA PRUDENTA (VALOAREA DE GARANTIE) — CRR III" in text
    # cifre auditabile: 280000 (piata/finala) -> 252000 (prudenta, -10%)
    assert "280000" in text and "252000" in text


def test_narrative_section():
    s = NarrativeSection(capitol="Analiza pietei", text="Piata locala...")
    assert s.capitol == "Analiza pietei"
    assert s.text == "Piata locala..."
