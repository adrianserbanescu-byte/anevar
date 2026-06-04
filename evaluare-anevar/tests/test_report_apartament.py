from decimal import Decimal

from docx import Document

from evaluare.models.meta import EvaluationMeta
from evaluare.models.narrative import NarrativeSection
from evaluare.models.property import BuildingData, CostElement, DepreciationPoint, LandData
from evaluare.models.report_context import ReportContext
from evaluare.models.results import CostResult, ReconciledResult
from evaluare.report.generator import genereaza_raport


def _ctx_base() -> ReportContext:
    meta = EvaluationMeta(
        client_nume="Ion Popescu", adresa="Str. Exemplu 1, Bucuresti",
        numar_cadastral="123456", carte_funciara="CF123456",
        evaluator_nume="Maria Ionescu", evaluator_legitimatie="19567",
        data_evaluarii="2026-01-16", data_raportului="2026-01-16",
    )
    return ReportContext(
        meta=meta,
        land=LandData(suprafata=Decimal("500"), categorie="intravilan"),
        building=BuildingData(
            au=Decimal("322.75"), acd=Decimal("351.46"), an_referinta=2025,
            elements=[
                CostElement(element="Structura", cod="X", um="mp",
                            cantitate=Decimal("100"), cost_unitar=Decimal("2000"),
                            an_pif=2015),
            ],
            depreciation_points=[DepreciationPoint(varsta=10, depreciere=Decimal("0.1"))],
        ),
        cost_result=CostResult(cib=Decimal("200000"), vcp=Decimal("10"),
                               depreciere_fizica=Decimal("0.10"), cin=Decimal("180000"),
                               valoare_cost=Decimal("280000")),
        reconciled=ReconciledResult(valoare_finala=Decimal("280000"),
                                    metoda_selectata="cost"),
        narrative=[NarrativeSection(capitol="Analiza celei mai bune utilizari (CMBU)",
                                    text="Cea mai buna utilizare este cea rezidentiala.")],
    )


def _all_text(path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_apartament_linie_prezenta_cand_etaj_setat(tmp_path):
    ctx = _ctx_base()
    ctx.building.etaj = 3
    ctx.building.nr_niveluri_bloc = 10
    ctx.building.an_bloc = 2008
    out = tmp_path / "raport_apt.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Apartament" in text
    assert "etaj 3" in text


def test_casa_fara_linie_apartament(tmp_path):
    ctx = _ctx_base()
    # building are etaj=None, nr_niveluri_bloc=None, an_bloc=None, cota_teren_indiviza=None
    out = tmp_path / "raport_casa.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Apartament:" not in text
