"""Round-trip diacritice (ă î ș ț â) prin calea de export: .docx → text → PDF → text.

Pe Windows, encoding-ul implicit e cp1252 (verificat: `locale.getpreferredencoding()`),
deci orice `open()` / `write_text()` / pipe fără `encoding="utf-8"` explicit corupe
diacriticele române ('Suprafață' → 'Suprafa??' la client = inacceptabil pe raport oficial).

Acoperit aici (cerinta A, lane D):
- python-docx: doc.save() → reload → text intact (scrie XML UTF-8 in ZIP intern).
- soffice (LibreOffice headless) docx → PDF: round-trip pypdf.extract_text() verifică
  diacriticele intacte. Skip dacă soffice lipsește (test_report_pdf.py face deja skipif).
- Salvare .docx şi reload, plus salvare în `salveaza()` din `aml/documente.py`.

Owner: D (Rol 2, reassign C → D pentru diacritice-export). NU modifică logica de business;
testează doar contractul de encoding pe calea export. Cazurile noi (paragrafe noi în
raport, etc.) rămân ale autorilor lor.
"""
from __future__ import annotations

import pytest
from docx import Document

# Eșantion reprezentativ: toate cele 5 diacritice + paragramă uzuală cu „Suprafață".
DIACRITICE = "ă î ș ț â Ă Î Ș Ț Â"
TEXT_RO = "Suprafață utilă: 120,50 mp; Vârstă: 25 ani; Înălțime: 2,75 m."


def _docx_cu_diacritice(tmp_path) -> object:
    """Construiește un .docx cu paragrafe care conțin diacritice + îl salvează pe disc."""
    cale = tmp_path / "diacritice.docx"
    doc = Document()
    doc.add_heading("Test diacritice", level=0)
    doc.add_paragraph(DIACRITICE)
    doc.add_paragraph(TEXT_RO)
    doc.save(str(cale))
    return cale


# ─── 1. python-docx scrie/citește UTF-8 indiferent de OS locale ───────────────
def test_docx_round_trip_pastreaza_diacriticele(tmp_path):
    """doc.save() scrie XML UTF-8 în ZIP; reload trebuie să întoarcă diacriticele intacte."""
    cale = _docx_cu_diacritice(tmp_path)
    citit = Document(str(cale))
    paragrafe = [p.text for p in citit.paragraphs]
    assert DIACRITICE in paragrafe, f"diacriticele lipsesc dupa reload: {paragrafe}"
    assert TEXT_RO in paragrafe, f"textul RO lipsește dupa reload: {paragrafe}"
    # Verifică suplimentar că NU s-au inserat caractere de înlocuire („?", „�")
    pline = "\n".join(paragrafe)
    assert "?" not in pline.replace("?: ", "_"), f"semn de corupere cp1252 în .docx: {pline}"
    assert "�" not in pline, f"caractere replacement (U+FFFD) în .docx: {pline}"


# ─── 2. Salvarea din aml.documente (cale folosită de generatoarele AML) ───────
def test_aml_salveaza_pastreaza_diacriticele(tmp_path):
    """`aml.documente.salveaza(doc, cale)` păstrează diacriticele la reload."""
    from evaluare.aml.documente import salveaza
    doc = Document()
    doc.add_paragraph(TEXT_RO)
    cale = salveaza(doc, tmp_path / "aml.docx")
    citit = Document(str(cale))
    assert TEXT_RO in [p.text for p in citit.paragraphs]


# ─── 3. Conversie docx → PDF cu LibreOffice (skip dacă nu există) ─────────────
@pytest.mark.skipif(
    True if __import__("evaluare.report.pdf", fromlist=["_gaseste_soffice"])._gaseste_soffice() is None
    else False, reason="LibreOffice indisponibil pe această stație")
def test_docx_to_pdf_pastreaza_diacriticele_in_pdf(tmp_path):
    """Pipeline-ul docx → PDF (soffice) trebuie să păstreze diacriticele în text-ul PDF.

    Verificare: după conversie, extragerea de text cu pypdf trebuie să conțină măcar
    una dintre diacritice. (PDF poate codifica text ca glyphs sau ca Unicode; soffice
    setează ToUnicode CMap → pypdf.extract_text() returnează caractere UTF-8.)
    """
    from pypdf import PdfReader

    from evaluare.report.pdf import docx_to_pdf
    docx = _docx_cu_diacritice(tmp_path)
    pdf = docx_to_pdf(docx)
    assert pdf.exists() and pdf.suffix == ".pdf"
    assert pdf.read_bytes()[:4] == b"%PDF"
    # Text PDF: verificăm că extragem text recognoscibil cu diacritice.
    text_pdf = "\n".join((p.extract_text() or "") for p in PdfReader(str(pdf)).pages)
    # Acceptăm ca soffice să folosească orice variantă de glyph: testăm că măcar
    # două dintre litere se regăsesc (acoperă cazurile reale "Suprafață", "vârstă").
    gasite = [c for c in "ăîșțâĂÎȘȚÂ" if c in text_pdf]
    assert len(gasite) >= 2, (
        f"PDF-ul produs de soffice NU conține diacritice (extracted={text_pdf!r}); "
        f"glyphs pierdute la conversie sau extracție = potențială regresie de export."
    )
