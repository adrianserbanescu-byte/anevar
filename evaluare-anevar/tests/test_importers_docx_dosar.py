"""Import dosar dintr-un `.docx`: parsarea numelui de fișier + extragerea din text."""
from __future__ import annotations

import docx

from evaluare.importers.docx_dosar import (
    extrage_din_docx,
    extrage_din_text,
    parse_nume_fisier,
)


# ── Numele fișierului (identitatea sigură) ───────────────────────────────────
def test_parse_nume_standard():
    d = parse_nume_fisier("21766 Bololoi Daniela-Doina locuinta Busteni.docx")
    assert d["id_client"] == "21766"
    assert d["nume_client"] == "Bololoi Daniela-Doina"
    assert d["tip_proprietate"] == "casa"          # „locuinta" -> casa
    assert d["localitate"] == "Busteni"
    assert d["judet"] == "Prahova"


def test_parse_nume_cu_strada_si_oras_dupa_virgula():
    d = parse_nume_fisier("21922 Cook Elena locuinta str.Fantanii ,Brasov.docx")
    assert d["nume_client"] == "Cook Elena"
    assert d["localitate"] == "Brasov"
    assert d["judet"] == "Brașov"


def test_parse_nume_localitate_compusa():
    d = parse_nume_fisier("21468 Paduraru Iulia-Gabriela casa Maneciu -Pamanteni.docx")
    assert d["tip_proprietate"] == "casa"
    assert d["localitate"] == "Maneciu-Pamanteni"
    assert d["judet"] == "Prahova"


def test_parse_nume_apartament():
    d = parse_nume_fisier("301 Ion Pop apartament Cluj-Napoca.docx")
    assert d["tip_proprietate"] == "apartament"
    assert d["judet"] == "Cluj"


def test_parse_nume_fara_id():
    d = parse_nume_fisier("Client fara numar teren Comarnic.docx")
    assert "id_client" not in d
    assert d["tip_proprietate"] == "agricol"        # „teren" -> agricol


# ── Textul raportului (best-effort) ──────────────────────────────────────────
def test_extrage_text_beneficiar_scop_data():
    text = ("Utilizarea desemnata a evaluării este garantarea împrumutului.\n"
            "se adreseaza ... UNICREDIT BANK SA in calitate de utilizator.\n"
            "Inspecția a fost efectuată la data de 20.03.2026 de evaluator.")
    d = extrage_din_text(text)
    assert d["scop"] == "garantare"
    assert d["beneficiar"] == "UNICREDIT BANK SA"
    assert d["data_inspectiei"] == "2026-03-20"


def test_extrage_text_raportare_financiara():
    d = extrage_din_text("Scopul este raportare financiară conform IFRS.")
    assert d["scop"] == "raportare_financiara"


def test_extrage_text_gol():
    assert extrage_din_text("text fără semnale relevante") == {}


# ── Integrare: docx sintetic (fișierele reale sunt private) ──────────────────
def test_extrage_din_docx_combina_filename_si_text(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Utilizarea desemnata a evaluării este garantarea împrumutului.")
    doc.add_paragraph("Clientul ION POPESCU, BANCA TRANSILVANIA BANK SA utilizator.")
    doc.add_paragraph("Inspecția la data de 04.04.2026 de evaluator.")
    f = tmp_path / "999 Ion Popescu casa Sinaia.docx"
    doc.save(str(f))
    w = extrage_din_docx(f)
    assert w["id_client"] == "999"
    assert w["nume_client"] == "Ion Popescu"
    assert w["tip_proprietate"] == "casa"
    assert w["judet"] == "Prahova"
    assert w["scop"] == "garantare"
    assert w["data_inspectiei"] == "2026-04-04"


def test_extrage_din_docx_fisier_invalid_ramane_pe_filename(tmp_path):
    f = tmp_path / "777 Maria Ionescu apartament Brasov.docx"
    f.write_text("not a real docx", encoding="utf-8")
    w = extrage_din_docx(f)                              # docx ilizibil -> doar filename
    assert w["id_client"] == "777"
    assert w["tip_proprietate"] == "apartament"
    assert w["scop"] == "garantare"                     # default
