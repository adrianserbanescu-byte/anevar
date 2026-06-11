"""Structura livrabilului (raportului) PER TIP de imobil.

Verifica tailoring-ul din generator.py condus de `ctx.profil.tip_activ` + `abordari_aplicabile`
(vezi profil.SECTIUNI_PER_TIP + docs/SEV-2025-cerinte-per-tip-imobil.md):
- teren / agricol: FARA sectiunea de constructie / cost de inlocuire constructie; analiza terenului ramane;
- apartament: FARA teren standalone + nota „cota parte indiviza din teren"; piata = principala;
- comercial: abordarea prin VENIT proeminenta / principala;
- casa: piata + cost.
Toate aditiv / backward-compatible (tipul implicit/necunoscut -> comportamentul actual).
"""
from decimal import Decimal

from docx import Document

from evaluare.engine.venit import DateVenit, RezultatVenit
from evaluare.models.comparable import Adjustment, Comparable, LandComparable
from evaluare.models.meta import EvaluationMeta
from evaluare.models.narrative import NarrativeSection
from evaluare.models.property import BuildingData, CostElement, DepreciationPoint, LandData
from evaluare.models.report_context import ReportContext
from evaluare.models.results import CostResult, MarketResult, ReconciledResult
from evaluare.profil import (
    AGRICOL,
    APARTAMENT_GARANTARE,
    CASA_TEREN_GARANTARE,
    COMERCIAL_INCHIRIAT,
    ProfilEvaluare,
)
from evaluare.report.generator import genereaza_raport


def _ctx(profil: ProfilEvaluare | None = None) -> ReportContext:
    meta = EvaluationMeta(
        client_nume="Ion Popescu", adresa="Str. Exemplu 1, Bucuresti",
        numar_cadastral="123456", carte_funciara="CF123456",
        evaluator_nume="Maria Ionescu", evaluator_legitimatie="19567",
        data_evaluarii="2026-01-16", data_raportului="2026-01-16",
    )
    ctx = ReportContext(
        meta=meta,
        land=LandData(suprafata=Decimal("500"), categorie="intravilan"),
        building=BuildingData(
            au=Decimal("322.75"), acd=Decimal("351.46"), an_referinta=2025,
            elements=[
                CostElement(element="Structura", cod="X", um="mp",
                            cantitate=Decimal("100"), cost_unitar=Decimal("2000"),
                            an_pif=2015),
            ],
            depreciation_points=[DepreciationPoint(varsta=10, depreciere=Decimal("0.1"))],
        ),
        cost_result=CostResult(cib=Decimal("200000"), vcp=Decimal("10"),
                               depreciere_fizica=Decimal("0.10"), cin=Decimal("180000"),
                               valoare_cost=Decimal("280000")),
        reconciled=ReconciledResult(valoare_finala=Decimal("280000"),
                                    metoda_selectata="cost"),
        narrative=[NarrativeSection(capitol="Analiza celei mai bune utilizari (CMBU)",
                                    text="Cea mai buna utilizare este cea rezidentiala.")],
    )
    if profil is not None:
        ctx.profil = profil
    return ctx


def _adauga_piata(ctx: ReportContext) -> None:
    ctx.comparables = [Comparable(pret=Decimal("300000"), suprafata=Decimal("320")),
                       Comparable(pret=Decimal("310000"), suprafata=Decimal("330"))]
    ctx.market_result = MarketResult(
        preturi_unitare_corectate=[Decimal("305000"), Decimal("308000")],
        ajustari_brute=[Decimal("0.05"), Decimal("0.08")],
        ajustari_nete=[Decimal("0.02"), Decimal("0.03")],
        index_selectat=0, valoare_piata=Decimal("305000"))


def _adauga_teren(ctx: ReportContext) -> None:
    from evaluare.engine.land import evaluate_land
    from evaluare.engine.reconciliation import aloca_constructii
    comps = [
        LandComparable(pret_mp=Decimal("100"), suprafata=Decimal("450"), localizare="Zona A",
                       adjustments=[Adjustment(element="A", tip="procentuala",
                                               valoare=Decimal("0.05"))]),
        LandComparable(pret_mp=Decimal("120"), suprafata=Decimal("500"),
                       adjustments=[Adjustment(element="A", tip="procentuala",
                                               valoare=Decimal("-0.30"))]),
    ]
    lr = evaluate_land(comps, ctx.land.suprafata)
    ctx.land_comparables = comps
    ctx.land_result = lr
    ctx.alocare_constructii = aloca_constructii(ctx.reconciled.valoare_finala, lr.valoare_teren)


def _adauga_venit(ctx: ReportContext) -> None:
    ctx.date_venit = DateVenit(venit_brut_potential=Decimal("100000"), grad_neocupare=Decimal("0.05"),
                               cheltuieli_exploatare=Decimal("20000"), rata_capitalizare=Decimal("0.08"))
    ctx.venit_result = RezultatVenit(noi=Decimal("76000.00"), valoare=Decimal("950000.00"))


def _all_text(path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# --------------------------------------------------------------------------- #
# TEREN / AGRICOL — fara sectiune de constructie, analiza terenului ramane
# --------------------------------------------------------------------------- #
def test_teren_nu_contine_sectiune_constructie(tmp_path):
    # Tip teren liber: chiar daca exista elemente de cost in input, raportul NU contine tabelul de
    # cost de inlocuire constructie, nici linia descriptiva a constructiei. Terenul ramane.
    ctx = _ctx(ProfilEvaluare(tip_activ="teren", abordari_aplicabile=["comparatie"], ghid="GEV_630"))
    _adauga_teren(ctx)
    out = tmp_path / "teren.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    # FARA constructie
    assert "Tabelul costurilor segregate" not in text
    assert "Constructie: Au" not in text
    assert "valoare prin cost" not in text
    # analiza terenului PASTRATA
    assert "Grila de comparatie pentru teren" in text
    assert "Teren: 500 mp" in text


def test_agricol_nu_contine_sectiune_constructie(tmp_path):
    # Tip agricol: la fel ca terenul — fara constructie; comparatia de teren ramane.
    ctx = _ctx(AGRICOL)
    ctx.land.categorie_folosinta = "arabil"
    ctx.land.clasa_calitate = 2
    _adauga_teren(ctx)
    out = tmp_path / "agricol.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Tabelul costurilor segregate" not in text
    assert "Constructie: Au" not in text
    assert "Teren agricol: categorie de folosinta arabil" in text   # descrierea terenului ramane


# --------------------------------------------------------------------------- #
# APARTAMENT — fara teren standalone + nota cota indiviza; piata principala
# --------------------------------------------------------------------------- #
def test_apartament_are_nota_cota_indiviza_fara_teren_standalone(tmp_path):
    # Apartament: terenul = cota indiviza netranzacționabila -> NU se evalueaza standalone si NU se
    # face alocare teren/constructie; in loc apare nota explicativa (GEV 630 §118.a).
    ctx = _ctx(APARTAMENT_GARANTARE)
    ctx.building.etaj = 3
    ctx.building.nr_niveluri_bloc = 10
    ctx.building.cota_teren_indiviza = Decimal("25")
    _adauga_piata(ctx)
    _adauga_teren(ctx)   # chiar daca avem grila de teren, apartamentul o suprima
    out = tmp_path / "apt.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    # nota cota indiviza prezenta
    assert "COTA-PARTE INDIVIZA" in text
    assert "GEV 630 §118.a" in text
    assert "Cota-parte indiviza aferenta: 25 mp" in text
    # FARA grila de teren standalone si FARA alocarea valorii
    assert "Grila de comparatie pentru teren" not in text
    assert "ALOCAREA VALORII" not in text
    # piata principala -> grila de piata prezenta
    assert "Grila de comparatie directa pe pret total" in text


# --------------------------------------------------------------------------- #
# COMERCIAL — venit ca abordare PRINCIPALA
# --------------------------------------------------------------------------- #
def test_comercial_venit_principal(tmp_path):
    # Comercial: abordarea prin VENIT este principala / proeminenta (GEV 232 §11).
    ctx = _ctx(COMERCIAL_INCHIRIAT)
    _adauga_piata(ctx)
    _adauga_venit(ctx)
    out = tmp_path / "comercial.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Abordarea PRINCIPALA pentru acest tip de proprietate (generatoare de venit)" in text
    assert "abordare PRINCIPALA" in text          # titlul sectiunii de venit marcat
    assert "Venit net din exploatare (NOI)" in text


def test_comercial_omite_cost_neaplicabil(tmp_path):
    # Comercial: costul de regula NU se aplica (SectiuniTip.abordare_cost=False) -> tabelul de cost NU
    # apare, chiar daca exista elemente de cost in input (din _ctx). GEV 232 §11.
    ctx = _ctx(COMERCIAL_INCHIRIAT)
    assert ctx.building.elements   # avem elemente de cost in input, dar tipul le suprima
    _adauga_venit(ctx)
    out = tmp_path / "comercial_fara_cost.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Tabelul costurilor segregate" not in text


# --------------------------------------------------------------------------- #
# CASA — piata + cost (comportamentul actual, backward-compatible)
# --------------------------------------------------------------------------- #
def test_casa_contine_cost_si_piata(tmp_path):
    # Casa: piata + cost — toate sectiunile clasice prezente (comportamentul actual).
    ctx = _ctx(CASA_TEREN_GARANTARE)
    _adauga_piata(ctx)
    _adauga_teren(ctx)
    out = tmp_path / "casa.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Tabelul costurilor segregate" in text                 # cost
    assert "Grila de comparatie directa pe pret total" in text    # piata
    assert "Grila de comparatie pentru teren" in text             # teren standalone
    assert "ALOCAREA VALORII" in text                             # alocare teren/constructie
    assert "Constructie: Au" in text
    assert "COTA-PARTE INDIVIZA" not in text                      # casa NU are nota de cota indiviza


def test_tip_implicit_pastreaza_comportamentul_actual(tmp_path):
    # Backward-compatible: fara profil per-tip explicit (profil implicit = casa), raportul iese
    # cu toate sectiunile data-driven, ca azi.
    ctx = _ctx()   # profil implicit (CASA_TEREN_GARANTARE)
    _adauga_piata(ctx)
    _adauga_teren(ctx)
    out = tmp_path / "implicit.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Tabelul costurilor segregate" in text
    assert "Grila de comparatie pentru teren" in text
    assert "ALOCAREA VALORII" in text
    assert "COTA-PARTE INDIVIZA" not in text
