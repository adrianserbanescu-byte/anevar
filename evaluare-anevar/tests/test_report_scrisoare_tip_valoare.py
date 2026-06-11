"""CONF-01 — scrisoarea de transmitere foloseste TIPUL valorii din profil, nu hardcodat „valoare de piata".

Pe profilul de garantare/piata textul ramane „valoarea de piata" (backward-compatible); pe profiluri
unde tipul valorii NU e valoarea de piata (asigurare = cost de reconstructie / SEV 450; raportare
financiara = valoare justa; lichidare etc.) scrisoarea afiseaza tipul corect, evitand contradictia
interna a raportului (restul raportului eticheteaza deja corect tipul valorii) si riscul la verificarea
bancara.
"""
from docx import Document

from evaluare.profil import ASIGURARE, RAPORTARE_FINANCIARA, ProfilEvaluare
from evaluare.report.generator import _tip_valoare_eticheta, genereaza_raport
from tests.test_report_generator import _ctx


def _text(path) -> str:
    doc = Document(str(path))
    parti = [p.text for p in doc.paragraphs]
    for tab in doc.tables:
        for row in tab.rows:
            parti.extend(c.text for c in row.cells)
    return "\n".join(parti)


def _scrisoare(text: str) -> str:
    """Doar paragraful scrisorii de transmitere (intre titlu si declaratia de conformitate)."""
    start = text.find("SCRISOARE DE TRANSMITERE")
    end = text.find("Raportul a fost elaborat in conformitate", start)
    assert start != -1
    return text[start:end if end != -1 else None]


def test_default_garantare_ramane_valoare_de_piata(tmp_path):
    # Backward-compat: profilul implicit (garantare / tip_valoare="piata") -> „valoarea de piata".
    out = tmp_path / "garantare.docx"
    genereaza_raport(_ctx(), out)
    scrisoare = _scrisoare(_text(out))
    assert "valoarea de piata estimata a proprietatii" in scrisoare


def test_asigurare_nu_spune_valoare_de_piata(tmp_path):
    # Profil ASIGURARE (SEV 450, tip_valoare="asigurare"): scrisoarea NU mai afirma „valoarea de piata".
    ctx = _ctx()
    ctx.profil = ASIGURARE
    out = tmp_path / "asigurare.docx"
    genereaza_raport(ctx, out)
    scrisoare = _scrisoare(_text(out))
    assert "valoarea de asigurare" in scrisoare
    assert "valoarea de piata estimata a proprietatii" not in scrisoare


def test_raportare_financiara_valoare_justa(tmp_path):
    # Profil RAPORTARE_FINANCIARA (tip_valoare="justa"): scrisoarea afirma „valoarea justa".
    ctx = _ctx()
    ctx.profil = RAPORTARE_FINANCIARA
    out = tmp_path / "justa.docx"
    genereaza_raport(ctx, out)
    scrisoare = _scrisoare(_text(out))
    assert "valoarea justa estimata a proprietatii" in scrisoare
    assert "valoarea de piata estimata a proprietatii" not in scrisoare


def test_lichidare_valoare_de_lichidare(tmp_path):
    # Profil cu tip_valoare="lichidare": scrisoarea afirma „valoarea de lichidare".
    ctx = _ctx()
    ctx.profil = ProfilEvaluare(tip_valoare="lichidare", abordari_aplicabile=["comparatie"])
    out = tmp_path / "lichidare.docx"
    genereaza_raport(ctx, out)
    scrisoare = _scrisoare(_text(out))
    assert "valoarea de lichidare estimata a proprietatii" in scrisoare
    assert "valoarea de piata estimata a proprietatii" not in scrisoare


def test_eticheta_fallback_pe_slug_necunoscut():
    # Slug necunoscut / profil fara tip_valoare valid -> „valoarea de piata" (comportamentul istoric).
    ctx = _ctx()
    ctx.profil = ProfilEvaluare()
    # Fortam un slug necunoscut pe instanta (model permisiv la atribut runtime) pentru a verifica fallback-ul.
    object.__setattr__(ctx.profil, "tip_valoare", "necunoscut")
    assert _tip_valoare_eticheta(ctx) == "valoarea de piata"
