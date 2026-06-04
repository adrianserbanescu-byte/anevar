from decimal import Decimal

from docx import Document

from evaluare.models.meta import EvaluationMeta
from evaluare.models.narrative import NarrativeSection
from evaluare.models.property import BuildingData, CostElement, DepreciationPoint, LandData
from evaluare.models.report_context import ReportContext
from evaluare.models.results import CostResult, ReconciledResult
from evaluare.profil import INDUSTRIAL
from evaluare.report.generator import genereaza_raport


def test_profil_industrial():
    assert INDUSTRIAL.tip_activ == "industrial"
    assert INDUSTRIAL.ghid == "GEV_630"
    assert INDUSTRIAL.abordari_aplicabile == ["cost", "venit", "comparatie"]


def test_building_inaltime_libera_optional():
    b = BuildingData(au=Decimal("300"), acd=Decimal("320"), an_referinta=2025,
                     inaltime_libera=Decimal("8"))
    assert b.inaltime_libera == Decimal("8")
    b2 = BuildingData(au=Decimal("300"), acd=Decimal("320"), an_referinta=2025)
    assert b2.inaltime_libera is None


def _ctx_industrial() -> ReportContext:
    meta = EvaluationMeta(
        client_nume="Ion Popescu", adresa="Str. Exemplu 1, Bucuresti",
        numar_cadastral="123456", carte_funciara="CF123456",
        evaluator_nume="Maria Ionescu", evaluator_legitimatie="19567",
        data_evaluarii="2026-01-16", data_raportului="2026-01-16",
    )
    return ReportContext(
        meta=meta,
        land=LandData(suprafata=Decimal("500"), categorie="intravilan"),
        building=BuildingData(
            au=Decimal("322.75"), acd=Decimal("351.46"), an_referinta=2025,
            elements=[
                CostElement(element="Structura", cod="X", um="mp",
                            cantitate=Decimal("100"), cost_unitar=Decimal("2000"),
                            an_pif=2015),
            ],
            depreciation_points=[DepreciationPoint(varsta=10, depreciere=Decimal("0.1"))],
        ),
        cost_result=CostResult(cib=Decimal("200000"), vcp=Decimal("10"),
                               depreciere_fizica=Decimal("0.10"), cin=Decimal("180000"),
                               valoare_cost=Decimal("280000")),
        reconciled=ReconciledResult(valoare_finala=Decimal("280000"),
                                    metoda_selectata="cost"),
        narrative=[NarrativeSection(capitol="Analiza celei mai bune utilizari (CMBU)",
                                    text="Cea mai buna utilizare este cea industriala.")],
    )


def _all_text(path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_raport_cu_inaltime_libera(tmp_path):
    ctx = _ctx_industrial()
    ctx.building.inaltime_libera = Decimal("8")
    out = tmp_path / "raport_industrial.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "inaltime libera 8" in text


def test_raport_fara_inaltime_libera(tmp_path):
    ctx = _ctx_industrial()
    # inaltime_libera is None by default — no industrial line should appear
    out = tmp_path / "raport_no_industrial.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Spatiu industrial" not in text
