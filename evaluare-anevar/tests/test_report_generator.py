from decimal import Decimal

from docx import Document

from evaluare.models.meta import EvaluationMeta
from evaluare.models.narrative import NarrativeSection
from evaluare.models.property import BuildingData, CostElement, DepreciationPoint, LandData
from evaluare.models.report_context import ReportContext
from evaluare.models.results import CostResult, ReconciledResult
from evaluare.report.generator import genereaza_raport


def _ctx() -> ReportContext:
    meta = EvaluationMeta(
        client_nume="Ion Popescu", adresa="Str. Exemplu 1, Bucuresti",
        numar_cadastral="123456", carte_funciara="CF123456",
        evaluator_nume="Maria Ionescu", evaluator_legitimatie="19567",
        data_evaluarii="2026-01-16", data_raportului="2026-01-16",
    )
    return ReportContext(
        meta=meta,
        land=LandData(suprafata=Decimal("500"), categorie="intravilan"),
        building=BuildingData(
            au=Decimal("322.75"), acd=Decimal("351.46"), an_referinta=2025,
            elements=[
                CostElement(element="Structura", cod="X", um="mp",
                            cantitate=Decimal("100"), cost_unitar=Decimal("2000"),
                            an_pif=2015),
            ],
            depreciation_points=[DepreciationPoint(varsta=10, depreciere=Decimal("0.1"))],
        ),
        cost_result=CostResult(cib=Decimal("200000"), vcp=Decimal("10"),
                               depreciere_fizica=Decimal("0.10"), cin=Decimal("180000"),
                               valoare_cost=Decimal("280000")),
        reconciled=ReconciledResult(valoare_finala=Decimal("280000"),
                                    metoda_selectata="cost"),
        narrative=[NarrativeSection(capitol="Analiza celei mai bune utilizari (CMBU)",
                                    text="Cea mai buna utilizare este cea rezidentiala.")],
    )


def _all_text(path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_gev520_inregistrare_big_conditionata_de_utilizator_desemnat(tmp_path):
    # GEV 520 ed. 2025, par. 77-78: la utilizator desemnat ANAF (garantare in reesalonarea datoriilor)
    # raportul NU se inregistreaza in BIG. Codul afirma acum CONDITIONAT, nu neconditionat. (Cross-check E.)
    # creditor (default): se inregistreaza in BIG
    out_c = tmp_path / "creditor.docx"
    genereaza_raport(_ctx(), out_c)
    tc = _all_text(out_c)
    assert "se inregistreaza in Baza Imobiliara de Garantii" in tc
    assert "NU se inregistreaza" not in tc
    assert "Raportul se inregistreaza in BIG." in tc                  # punct de checklist (creditor)
    # ANAF: NU se inregistreaza in BIG
    ctx = _ctx()
    ctx.meta.utilizator_desemnat = "ANAF"
    out_a = tmp_path / "anaf.docx"
    genereaza_raport(ctx, out_a)
    ta = _all_text(out_a)
    assert "utilizatorul desemnat ANAF" in ta
    assert "NU se inregistreaza in Baza Imobiliara de Garantii" in ta
    assert "Raportul se inregistreaza in BIG." not in ta             # checklistul NU mai afirma inregistrarea
    assert "NU se inregistreaza in BIG (GEV 520 par. 77-78)" in ta   # punct de checklist (ANAF)


def test_genereaza_raport_creeaza_fisier(tmp_path):
    out = tmp_path / "raport.docx"
    genereaza_raport(_ctx(), out)
    assert out.exists()
    # se deschide ca document Word valid
    Document(str(out))


def test_raportul_contine_datele_cheie(tmp_path):
    out = tmp_path / "raport.docx"
    genereaza_raport(_ctx(), out)
    text = _all_text(out)
    assert "Ion Popescu" in text
    assert "Garantarea creditului ipotecar" in text
    assert "280.000" in text                       # valoarea finala (formatata)
    assert "CIN" in text or "180.000" in text       # tabel cost
    assert "Cea mai buna utilizare" in text        # narativ inserat


def test_raportul_contine_disclaimer_aplicatie(tmp_path):
    # Cerinta Adi: ORICE document generat poarta disclaimer-ul aplicatiei catre evaluator (draft +
    # raspunderea revine evaluatorului semnatar + aplicatia fara raspundere) — mereu, nu doar in demo.
    from evaluare.report.generator import DISCLAIMER_APLICATIE
    out = tmp_path / "raport.docx"
    genereaza_raport(_ctx(), out)              # adnotari=False -> raport normal, nu demo
    text = _all_text(out)
    assert "NOTĂ A APLICAȚIEI CĂTRE EVALUATOR" in text
    assert DISCLAIMER_APLICATIE in text
    assert "DRAFT" in text and "nicio răspundere" in text


def test_raport_offline_fara_ai_e_complet(tmp_path):
    # ADR-004 (offline-first): fără narativ AI (offline / fără client), raportul iese COMPLET
    # cu text-șablon (fallback), fără să eșueze — dependența online e DOAR pentru lustruirea AI.
    ctx = _ctx()
    ctx.narrative = []                 # nicio secțiune AI (ca atunci când nu există client/internet)
    out = tmp_path / "raport.docx"
    genereaza_raport(ctx, out)
    doc = Document(str(out))
    titluri = "\n".join(p.text for p in doc.paragraphs if p.style.name.startswith("Heading"))
    for nr in ["1.", "2.", "3.", "4.", "5.", "6.", "7."]:
        assert nr in titluri           # toate cele 7 capitole prezente
    assert "[de completat]" in _all_text(out)   # fallback de șablon pentru secțiunile narative


def test_raport_sectiuni_rare_piata_venit_dcf_apartament(tmp_path):
    # Acoperă secțiunile rare: grila de comparație piață + abordarea prin venit + DCF + detalii apartament.
    from evaluare.engine.venit import DateVenit, RezultatVenit
    from evaluare.models.comparable import Comparable
    from evaluare.models.results import MarketResult
    ctx = _ctx()
    ctx.building.etaj = 3                          # apartament -> acoperă blocul 682-690
    ctx.building.nr_niveluri_bloc = 8
    ctx.building.an_bloc = 1985
    ctx.building.cota_teren_indiviza = Decimal("25")
    ctx.building.inaltime_libera = Decimal("4.5")  # spațiu industrial -> acoperă linia 692
    ctx.comparables = [Comparable(pret=Decimal("300000"), suprafata=Decimal("320")),
                       Comparable(pret=Decimal("310000"), suprafata=Decimal("330"))]
    ctx.market_result = MarketResult(
        preturi_unitare_corectate=[Decimal("305000"), Decimal("308000")],
        ajustari_brute=[Decimal("0.05"), Decimal("0.08")],
        ajustari_nete=[Decimal("0.02"), Decimal("0.03")],
        index_selectat=0, valoare_piata=Decimal("305000"))
    ctx.date_venit = DateVenit(venit_brut_potential=Decimal("36000"), grad_neocupare=Decimal("0.1"),
                               cheltuieli_exploatare=Decimal("6000"), rata_capitalizare=Decimal("0.08"))
    ctx.venit_result = RezultatVenit(noi=Decimal("26400"), valoare=Decimal("330000"))
    ctx.dcf_valoare = Decimal("325000")
    out = tmp_path / "raport.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Grila de comparatie directa pe pret total" in text       # grila piață (351-371)
    assert "Abordarea prin venit (capitalizare directă)" in text     # venit (726-737)
    assert "Abordarea prin venit (DCF)" in text                      # DCF (738-740)
    assert "etaj 3/8" in text and "an bloc 1985" in text             # detalii apartament
    assert "Spatiu industrial" in text                               # secțiune industrială (692)


def test_tip_valoare_txt_cazuri_limita():
    from evaluare.report.generator import _tip_valoare_txt
    assert "SEV 102" in _tip_valoare_txt("tip-necunoscut-oarecare")  # slug necunoscut -> ref. adăugată
    deja = "Valoare specială (SEV 104)"
    assert _tip_valoare_txt(deja) == deja                            # frază deja citată -> neschimbată
    assert _tip_valoare_txt("") is not None                          # gol -> nu crapă


def test_tipul_valorii_citeaza_sursa(tmp_path):
    # SEV 102 §20.4 + SEV 106 §30.6(i): slug-ul de tip valoare („piata", cum îl setează un profil)
    # devine denumire lizibilă + sursa definiției în raport (nu mai afișează slug-ul brut).
    ctx = _ctx()
    ctx.meta.tip_valoare = "piata"
    out = tmp_path / "raport.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Tipul valorii" in text
    assert "valoare de piață" in text                # slug -> denumire lizibilă
    assert "SEV 102" in text                          # sursa/definiția tipului valorii
    assert "estimate: piata." not in text             # nu mai afișează slug-ul brut


def test_raportul_declara_dreptul_si_sarcinile(tmp_path):
    # SEV 230 §40.1/§140: raportul declară dreptul evaluat + sarcinile CF (critic la garantare).
    ctx = _ctx()
    ctx.meta.sarcini = "ipotecă în favoarea Băncii X"
    out = tmp_path / "raport.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Dreptul evaluat" in text and "SEV 230" in text
    assert "ipotecă în favoarea Băncii X" in text


def test_sarcini_nedeclarate_avertizeaza(tmp_path):
    out = tmp_path / "raport.docx"
    genereaza_raport(_ctx(), out)            # sarcini=None implicit
    text = _all_text(out)
    assert "Sarcini" in text and "de verificat în extrasul de carte funciară" in text


def test_raportul_declara_inspectia(tmp_path):
    # GEV 630 §24/§44: raportul declară amploarea inspecției + însoțitorul.
    ctx = _ctx()
    ctx.meta.data_inspectiei = "2026-01-15"
    ctx.meta.inspectie_amploare = "interior și exterior"
    ctx.meta.inspectie_insotitor = "proprietarul"
    out = tmp_path / "raport.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Inspectia proprietatii" in text
    assert "interior și exterior" in text and "proprietarul" in text


def test_raportul_descrie_utilitati_si_urbanism(tmp_path):
    # GEV 630 §28/§16: utilitățile + regimul urbanistic apar în descrierea proprietății.
    ctx = _ctx()
    ctx.land.utilitati = ["apă", "energie electrică"]
    ctx.land.restrictii_urbanism = "POT 40%, CUT 1.2"
    ctx.land.acces = "drum asfaltat"
    out = tmp_path / "raport.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Utilitati:" in text and "energie electrică" in text
    assert "Regim urbanistic" in text and "POT 40%" in text
    assert "Acces:" in text and "drum asfaltat" in text


def test_raportul_are_cele_sapte_capitole(tmp_path):
    out = tmp_path / "raport.docx"
    genereaza_raport(_ctx(), out)
    doc = Document(str(out))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    titlu = "\n".join(headings)
    for nr in ["1.", "2.", "3.", "4.", "5.", "6.", "7."]:
        assert nr in titlu


def test_adnotari_demo_doar_cand_sunt_cerute(tmp_path):
    fara = tmp_path / "fara.docx"
    genereaza_raport(_ctx(), fara, adnotari=False)
    assert "NOTA DEMO" not in _all_text(fara)
    assert "LEGENDA NOTELOR DEMO" not in _all_text(fara)

    cu = tmp_path / "cu.docx"
    genereaza_raport(_ctx(), cu, adnotari=True)
    text = _all_text(cu)
    assert "NOTA DEMO" in text
    assert "LEGENDA NOTELOR DEMO" in text
    assert "[CALCULAT]" in text and "[AI]" in text and "[EXEMPLU]" in text


def test_termeni_referinta_acopera_sev_101_si_esg(tmp_path):
    out = tmp_path / "raport.docx"
    genereaza_raport(_ctx(), out)
    text = _all_text(out)
    assert "Sursa informatiilor" in text          # SEV 101 20.1 j
    assert "ESG" in text and "guvernanta" in text  # SEV 101/106 m (nou 2025)
    assert "Natura si amploarea activitatilor" in text  # 20.1 i
    assert "difuzarea sau publicarea" in text      # 20.1 o
    assert "Specialist" in text                    # 30.6 o


def test_raportul_eticheteaza_ai_vs_determinist(tmp_path):
    # Mitigare percepție bănci: raportul declară explicit că TOATE numerele sunt deterministe,
    # AI scrie doar proza, iar evaluatorul verifică și își asumă.
    out = tmp_path / "raport.docx"
    genereaza_raport(_ctx(), out)
    text = _all_text(out)
    assert "Asistare software si componenta AI" in text
    assert "DETERMINISTE" in text
    assert "NU sunt generate de inteligenta artificiala" in text
    assert "asumat integral de evaluatorul" in text
    assert "nu il inlocuieste si nu decide valoarea" in text


def test_raportul_conform_sev_2025_si_gev_520(tmp_path):
    out = tmp_path / "raport.docx"
    genereaza_raport(_ctx(), out)
    text = _all_text(out)
    # editia 2025 citata
    assert "editia 2025" in text or "ediția 2025" in text
    # tip valoare actualizat (SEV 102, nu SEV 104)
    assert "SEV 102" in text
    assert "SEV 104" not in text
    # factorii obligatorii GEV 520 A5
    assert "GEV 520, A5" in text
    assert "tendintele pietei" in text
    assert "valorii viitoare a garantiei" in text
    # inregistrare BIG
    assert "Baza Imobiliara de Garantii (BIG)" in text


def test_raportul_arata_echivalent_lei(tmp_path):
    ctx = _ctx()
    ctx.meta.moneda = "EUR"
    ctx.meta.curs_valutar = Decimal("5")     # 280000 EUR -> 1.400.000 LEI
    out = tmp_path / "raport.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "echivalent" in text
    assert "1.400.000 LEI" in text


def test_raportul_are_shell_gbf(tmp_path):
    out = tmp_path / "raport.docx"
    genereaza_raport(_ctx(), out)
    text = _all_text(out)
    for sectiune in ["SCRISOARE DE TRANSMITERE", "DECLARATIE DE CONFORMITATE",
                     "TERMENI DE REFERINTA", "RISCUL ASOCIAT GARANTIEI (GEV 520)",
                     "ANEXE", "Intocmit,"]:
        assert sectiune in text


def _ctx_cu_teren() -> ReportContext:
    from evaluare.engine.land import evaluate_land
    from evaluare.engine.reconciliation import aloca_constructii
    from evaluare.models.comparable import Adjustment, LandComparable
    ctx = _ctx()
    comps = [
        LandComparable(pret_mp=Decimal("100"), suprafata=Decimal("450"), localizare="Zona A",
                       adjustments=[Adjustment(element="A", tip="procentuala",
                                               valoare=Decimal("0.05"))]),
        LandComparable(pret_mp=Decimal("120"), suprafata=Decimal("500"),
                       adjustments=[Adjustment(element="A", tip="procentuala",
                                               valoare=Decimal("-0.30"))]),
    ]
    lr = evaluate_land(comps, ctx.land.suprafata)
    ctx.land_comparables = comps
    ctx.land_result = lr
    ctx.alocare_constructii = aloca_constructii(ctx.reconciled.valoare_finala, lr.valoare_teren)
    return ctx


_PNG_1x1 = ("data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ"
            "AAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg==")


def test_raportul_embeddeaza_fotografiile(tmp_path):
    ctx = _ctx()
    ctx.photos = [_PNG_1x1, "data:image/png;base64,GUNOI_INVALID==", _PNG_1x1]
    out = tmp_path / "raport.docx"
    genereaza_raport(ctx, out)
    doc = Document(str(out))
    # cele doua imagini valide sunt inserate; cea invalida e sarita fara eroare
    assert len(doc.inline_shapes) == 2
    text = _all_text(out)
    assert "Planse fotografice" in text
    assert "[de atasat]" not in text.split("Anexa 2")[1].split("Anexa 3")[0]


def test_raportul_embeddeaza_documentele_anexa3(tmp_path):
    ctx = _ctx()
    ctx.documente = [_PNG_1x1]
    out = tmp_path / "raport.docx"
    genereaza_raport(ctx, out)
    doc = Document(str(out))
    assert len(doc.inline_shapes) >= 1
    text = _all_text(out)
    dupa_anexa3 = text.split("Anexa 3")[1]
    assert "[de atasat]" not in dupa_anexa3


def test_raportul_anexa_foto_placeholder_fara_poze(tmp_path):
    out = tmp_path / "raport.docx"
    genereaza_raport(_ctx(), out)   # fara photos
    text = _all_text(out)
    dupa_anexa2 = text.split("Anexa 2")[1].split("Anexa 3")[0]
    assert "[de atasat]" in dupa_anexa2


def test_raportul_contine_grila_teren_si_alocare(tmp_path):
    out = tmp_path / "raport.docx"
    genereaza_raport(_ctx_cu_teren(), out)
    text = _all_text(out)
    assert "Grila de comparatie pentru teren" in text
    assert "Pret/mp corectat" in text
    assert "ALOCAREA VALORII" in text
    assert "Valoarea constructiilor" in text


def test_anexe_corupte_logheaza_avertisment_nu_dispar_tacut(tmp_path, caplog):
    """Eșec silențios reparat: o fotografie/scan care nu se poate insera lasă acum o urmă în jurnal
    (nu mai dispare tăcut din raport), iar raportul se generează oricum."""
    import base64
    import logging

    ctx = _ctx()
    # foto: base64 valid dar NU e imagine -> add_picture eșuează (ramura `except` reparată)
    ctx.photos = ["data:image/png;base64," + base64.b64encode(b"nu este o imagine reala").decode()]
    # document: base64 invalid -> _decode_foto întoarce None (ramura de decodare)
    ctx.documente = ["@@@base64-invalid@@@"]
    out = tmp_path / "raport.docx"
    with caplog.at_level(logging.WARNING, logger="evaluare.report.generator"):
        genereaza_raport(ctx, out)
    assert out.exists()                                   # raportul se generează oricum (nu crapă)
    assert "Anexa 2: fotografia 1 nu a putut fi inserată" in caplog.text
    assert "Anexa 3: documentul 1 nu a putut fi decodat" in caplog.text


def test_raport_afiseaza_nota_reconciliere(tmp_path):
    # Transparenta (optiunea b, decizia Adi): nota de reconciliere (ex. abordare calculata dar
    # neponderata) apare EXPLICIT in raport, ca valoarea finala sa nu diverge tacit de indicatii.
    ctx = _ctx()
    ctx.reconciled.nota = "Abordarea prin venit a fost calculata dar NU este inclusa in valoarea ponderata."
    out = tmp_path / "r.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Notă privind reconcilierea" in text and "venit" in text


def test_sev106_raport_contine_cele_18_elemente_obligatorii(tmp_path):
    # SEV 106 §30.6 (continutul minim al raportului de evaluare): test STRUCTURAL ca raportul contine
    # cele 18 elemente obligatorii. Test de completitudine — prinde regresii care scot vreun element.
    out = tmp_path / "r.docx"
    genereaza_raport(_ctx(), out)
    t = _all_text(out)
    tl = t.lower()
    elemente = {
        "1. evaluator (nume + legitimatie)": "Maria Ionescu" in t and "19567" in t,
        "2. client": "Ion Popescu" in t,
        "3. scopul evaluarii": "garantarea" in tl or "scop" in tl,
        "4. identificarea proprietatii (adresa)": "Str. Exemplu" in t,
        "5. nr. cadastral + carte funciara": "123456" in t and "CF123456" in t,
        "6. tipul valorii (SEV 102)": "SEV 102" in t or "valoare de piat" in tl,
        "7. data evaluarii si a raportului": "2026-01-16" in t,
        "8. ipoteze (generale si speciale)": "ipotez" in tl,
        "9. abordarile / metodele aplicate": "abordare" in tl or "metodelor" in tl,
        "10. reconcilierea rezultatelor": "reconcili" in tl,
        "11. valoarea finala estimata": "valoarea estimata" in tl or "valoarea finala" in tl,
        "12. declaratia de conformitate": "conformitate" in tl,
        "13. restrictii de utilizare / utilizator desemnat": (
            "exclusiv" in tl or "desemnat" in tl or "difuzar" in tl),
        "14. conformitatea cu standardele SEV/IVS": "SEV" in t,
        "15. dreptul de proprietate evaluat": "drept" in tl,
        "16. amploarea inspectiei / investigatiei": "inspect" in tl or "investiga" in tl,
        "17. cea mai buna utilizare (CMBU)": "cmbu" in tl or "cea mai buna utilizare" in tl,
        "18. precizarea fara TVA": "tva" in tl,
    }
    lipsa = [k for k, ok in elemente.items() if not ok]
    assert not lipsa, f"Elemente SEV 106 §30.6 lipsa din raport: {lipsa}"


def test_sev100_declara_scepticism_si_verificare_calitate(tmp_path):
    # SEV 100 (2025): raportul declara EXPLICIT scepticismul profesional (§10.4) + procedura de
    # verificare a calitatii (§20) — le aplicam de-facto (validari + audit), dar trebuie DECLARATE.
    out = tmp_path / "r.docx"
    genereaza_raport(_ctx(), out)
    t = _all_text(out)
    assert "scepticism profesional" in t and "SEV 100, par. 10.4" in t
    assert "verificare a calitatii" in t and "SEV 100, par. 20" in t


def test_sev450_asigurare_valoare_cib_brut_nu_cin_net(tmp_path):
    # E-(a) SEV 450: la scop=asigurare valoarea = cost de RECONSTRUCTIE (CIB brut, NEdeprecat, fara teren),
    # NU CIN net+teren (altfel sub-asigurare). Referinta = SEV 450, nu GEV 630.
    from decimal import Decimal

    from evaluare.assembler import EvaluationInput, construieste_context
    from evaluare.engine.cost import evaluate_cost
    building = {"au": "100", "acd": "120", "an_referinta": 2025,
                "elements": [{"element": "S", "cod": "X", "um": "mp", "cantitate": "120",
                              "cost_unitar": "2000", "an_pif": 2015}],
                "depreciation_points": [{"varsta": 5, "depreciere": "0.05"}]}
    inp = EvaluationInput(
        meta={"client_nume": "Ion", "adresa": "A", "numar_cadastral": "1", "carte_funciara": "CF",
              "evaluator_nume": "E", "evaluator_legitimatie": "1", "data_evaluarii": "2026-01-01",
              "data_raportului": "2026-01-01", "tip_valoare": "asigurare"},
        land={"suprafata": "500"}, building=building,
        valoare_teren="100000", metoda="cost", scop="asigurare")
    ctx = construieste_context(inp, client=None)
    cost = evaluate_cost(inp.building, valoare_teren=Decimal("100000"))
    assert ctx.reconciled.valoare_finala == cost.cib.quantize(Decimal("0.01"))   # CIB brut
    assert ctx.reconciled.valoare_finala != cost.valoare_cost                    # NU CIN net + teren
    out = tmp_path / "asig.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "SEV 450" in text and "reconstruc" in text.lower()
