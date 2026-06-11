"""UI nou (versiunea curentă): cont local + ÎNCEPE + workspace dosar (stocare pe foldere)."""
from __future__ import annotations

import pytest
from fastapi.testclient import TestClient

from evaluare.db.storage import Storage
from evaluare.web.app import create_app


def _fake_pdf(docx):
    """Convertor PDF fals (deterministic, fără LibreOffice) pentru teste."""
    from pathlib import Path
    p = Path(docx).with_suffix(".pdf")
    p.write_bytes(b"%PDF-1.4 fake\n%%EOF")
    return p


@pytest.fixture
def client(tmp_path, monkeypatch):
    monkeypatch.setenv("OUTPUT_DIR", str(tmp_path / "date"))
    s = Storage(tmp_path / "d.db")
    s.init()
    c = TestClient(create_app(storage=s, client=None, pdf_converter=_fake_pdf))
    c._baza = tmp_path
    c.storage = s
    return c


def _payload():
    return {
        "meta": {"client_nume": "Ion Pop", "adresa": "A", "numar_cadastral": "777",
                 "carte_funciara": "CF", "evaluator_nume": "Adi S", "evaluator_legitimatie": "8717",
                 "data_evaluarii": "2026-01-16", "data_raportului": "2026-01-16"},
        "land": {"suprafata": "500"},
        "building": {"au": "100", "acd": "120", "an_referinta": 2025,
                     "elements": [{"element": "S", "cod": "X", "um": "mp", "cantitate": "120",
                                   "cost_unitar": "2000", "an_pif": 2015}],
                     "depreciation_points": [{"varsta": 5, "depreciere": "0.05"}]},
        "valoare_teren": "100000", "metoda": "cost",
    }


def _cont(client):
    return client.post("/api/cont", json={
        "nume": "Adi S", "legitimatie": "8717",
        "format_dosar": ["id_client", "nume_client", "tip_proprietate"]})


# ── Cont ─────────────────────────────────────────────────────────────────────
def test_pagina_cont_fara_cont(client):
    r = client.get("/cont")
    assert r.status_code == 200 and "Nume evaluator" in r.text


def test_creeaza_cont_si_reincarca(client):
    assert _cont(client).status_code == 200
    r = client.get("/cont")
    assert "Adi S" in r.text and "8717" in r.text


def test_creare_dosar_nu_logheaza_numele_evaluatorului(client, caplog):
    """Igiena PII (audit #9, gap testare flag-at de D): la crearea unui dosar logam DOAR
    legitimatia (ID profesional), NICIODATA numele evaluatorului. Aserțiune POZITIVĂ + NEGATIVĂ."""
    import logging
    _cont(client)
    caplog.clear()                                   # ignoram logurile de la crearea contului
    with caplog.at_level(logging.INFO):
        uid = client.post("/api/dosar", json={"wizard": {"id_client": "D1"}}).json()["uuid"]
    assert uid                                       # dosar creat
    assert "8717" in caplog.text                     # POZITIV: legitimatia (ID prof) E logata
    assert "Adi S" not in caplog.text                # NEGATIV: numele evaluatorului NU apare in log


def test_cont_invalid_422(client):
    assert client.post("/api/cont", json={"nume": "", "legitimatie": "x"}).status_code == 422


# ── ÎNCEPE ───────────────────────────────────────────────────────────────────
def test_incepe_fara_cont_cere_cont(client):
    r = client.get("/incepe")
    assert r.status_code == 200 and "Creează cont" in r.text


def test_incepe_cu_cont_arata_optiuni(client):
    _cont(client)
    r = client.get("/incepe")
    assert "Dosar nou" in r.text and "Adi S" in r.text


# ── Dosar (workspace) ────────────────────────────────────────────────────────
def test_dosar_nou_necesita_cont(client):
    assert client.post("/api/dosar", json={"wizard": {}}).status_code == 403


def test_creeaza_deschide_salveaza_dosar(client):
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {"nume_client": "Ana"}}).json()["uuid"]
    assert client.get(f"/dosar/{uid}").status_code == 200
    # salvează snapshot
    client.post(f"/api/dosar/{uid}/salveaza", json={"nume_client": "Ana", "au": "90"})
    d = client.get(f"/api/dosar/{uid}").json()
    assert d["wizard"]["au"] == "90"


def test_dosar_apare_in_incepe(client):
    _cont(client)
    client.post("/api/dosar", json={"wizard": {"nume_client": "Vasile"}})
    client.get("/incepe")                                   # prima vedere -> intră în index
    assert "/dosar/" in client.get("/incepe").text


def test_calcul_fara_persistenta_sqlite(client):
    # «Calculează» din UI-ul nou NU trebuie să scrie rânduri orfane în SQLite (folder=adevăr)
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    # Numărăm direct în SQLite (sursa de adevăr); /dosare e acum redirect (gap UX F1-1).
    n_inainte = len(client.storage.list())
    r = client.post(f"/api/dosar/{uid}/calcul", json=_payload())
    assert r.status_code == 200
    d = r.json()
    assert d["valoare_finala"] and d["metoda"] == "cost"
    assert "id" not in d                                    # calc-only: nu persistă în SQLite
    assert len(client.storage.list()) == n_inainte          # zero dosare noi în SQLite
    assert client.post("/api/dosar/nope/calcul", json=_payload()).status_code == 404


def test_calcul_metoda_venit(client):
    # paritate UI nou: metoda venit (capitalizare directă) cu date_venit
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    p = _payload()
    p["metoda"] = "venit"
    p["date_venit"] = {"venit_brut_potential": "24000", "grad_neocupare": "0.1",
                       "cheltuieli_exploatare": "3000", "rata_capitalizare": "0.08"}
    r = client.post(f"/api/dosar/{uid}/calcul", json=p)
    assert r.status_code == 200 and r.json()["metoda"] == "venit"


def test_calcul_cu_grila_teren(client):
    # paritate UI nou: comparabile de teren → valoarea terenului se calculează din grilă
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    p = _payload()
    p["land_comparables"] = [{"pret_mp": "50", "suprafata": "700"},
                             {"pret_mp": "55", "suprafata": "650"},
                             {"pret_mp": "48", "suprafata": "720"}]
    r = client.post(f"/api/dosar/{uid}/calcul", json=p)
    assert r.status_code == 200 and r.json()["valoare_finala"]


def test_genereaza_raport_salveaza_versiune(client):
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    r = client.post(f"/api/dosar/{uid}/raport.docx", json=_payload())
    assert r.status_code == 200 and len(r.content) > 1000
    folder = client._baza / "date" / "dosare" / uid
    assert len(list(folder.glob("raport-*.docx"))) == 1


def test_genereaza_cu_adnotari(client):
    # opțiune Generează: adnotări de proveniență (raport de review)
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    r = client.post(f"/api/dosar/{uid}/raport.docx?adnotari=1", json=_payload())
    assert r.status_code == 200 and len(r.content) > 1000


def test_audit_txt_pe_flux_foldere(client):
    # Audit in-place: urma de audit pe fluxul nou (foldere), fără SQLite
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    r = client.post(f"/api/dosar/{uid}/audit.txt", json=_payload())
    assert r.status_code == 200 and "valoare_finala" in r.text
    assert client.post("/api/dosar/nope/audit.txt", json=_payload()).status_code == 404


def test_genereaza_raport_cu_anexa_foto(client):
    # Anexe: o fotografie (data-URL) ajunge în raport prin photos (Anexa 2)
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    png = ("data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAA"
           "C0lEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg==")
    p = _payload()
    p["photos"] = [png]
    r = client.post(f"/api/dosar/{uid}/raport.docx", json=p)
    assert r.status_code == 200 and len(r.content) > 1000


def test_backup_dosare_zip(client):
    # Backup: arhivează toate dosarele într-un .zip
    import io
    import zipfile
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {"nume_client": "Z"}}).json()["uuid"]
    r = client.get("/api/backup-dosare.zip")
    assert r.status_code == 200
    z = zipfile.ZipFile(io.BytesIO(r.content))
    assert any(uid in n and n.endswith("dosar.json") for n in z.namelist())


# ── Hardening din auditul final (securitate + buguri) ────────────────────────
def test_grila_casa_goala_422(client):
    # bug: <3 comparabile -> ValueError din motor; trebuie 422 clar, NU 500.
    assert client.post("/api/grila-casa", json={"suprafata_subiect": "120", "comparabile": []}).status_code == 422
    assert client.post("/api/grila-teren", json={"suprafata_subiect": "500", "comparabile": []}).status_code == 422


def test_raport_date_insuficiente_422(client):
    # robustețe: cost fără puncte de depreciere -> 422 clar (din motor), NU 500 nehandelat.
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    p = _payload()
    p["building"]["depreciation_points"] = []
    assert client.post(f"/api/dosar/{uid}/raport.docx", json=p).status_code == 422
    assert client.post(f"/api/dosar/{uid}/calcul", json=p).status_code == 422


def test_raport_blocat_pe_problema_blocanta_422(client):
    # I1 (re-audit lean — gap de testare la nivel de router): o problemă `nivel="blocheaza"`
    # (Au>Acd) e ADVISORY în /calcul (200 + alertă; evaluatorul vede valoarea+problema), dar
    # BLOCHEAZĂ documentul OFICIAL /raport.docx (422 „Raport blocat") — un raport SEMNABIL de
    # garantare NU se generează pe date blocante. Regresie pt commit 7d6ed0a.
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    p = _payload()
    p["building"]["au"] = "200"                       # Au(200) > Acd(120) -> nivel="blocheaza"
    # /calcul: calcul reușește, problema rămâne advisory (200 + alertă)
    rc = client.post(f"/api/dosar/{uid}/calcul", json=p)
    assert rc.status_code == 200 and rc.json()["alerte"]
    # /raport.docx: ACELAȘI input -> documentul oficial e BLOCAT (422), nu generat mut
    rr = client.post(f"/api/dosar/{uid}/raport.docx", json=p)
    assert rr.status_code == 422 and "blocat" in rr.json()["detail"].lower()


def test_validare_422_mesaj_citibil_nu_object_object(client):
    # M1 (audit user-journey): un câmp obligatoriu lipsă (numar_cadastral, schema-required) producea
    # un array Pydantic BRUT pe care UI-ul îl afișa evaluatorului ca „[object Object]". Handler-ul global
    # RequestValidationError formatează în `detail` STRING citibil (numește câmpul) -> mesaj prietenos.
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    p = _payload()
    del p["meta"]["numar_cadastral"]                  # câmp obligatoriu lipsă -> 422 Pydantic
    r = client.post(f"/api/dosar/{uid}/calcul", json=p)
    assert r.status_code == 422
    det = r.json()["detail"]
    assert isinstance(det, str)                       # STRING, NU array (array -> „[object Object]" în UI)
    assert "numar_cadastral" in det                   # numește câmpul-problemă, acționabil


def test_import_docx_continut_invalid_nu_crapa(client):
    # robustețe: conținut care NU e .docx valid -> NU 500; degradează grațios la parsarea numelui.
    import base64
    _cont(client)
    payload = base64.b64encode(b"acesta nu este un document Word").decode()
    r = client.post("/api/dosar/import-docx",
                    json={"nume_fisier": "1_Test_casa.docx", "continut": payload})
    assert r.status_code == 200 and "uuid" in r.json()


def test_evaluare_veche_date_insuficiente_422(client):
    # endpoint-ul vechi /api/evaluare nu trebuie să dea 500 pe date incomplete.
    _cont(client)
    p = _payload()
    p["building"]["depreciation_points"] = []
    assert client.post("/api/evaluare", json=p).status_code == 422


def test_raport_mereu_docx_fara_pdf(client):
    # Decizie 2026-06-08: app-ul NU mai produce PDF. Orice generare -> .docx;
    # un ?fmt=pdf de la un frontend vechi e IGNORAT (forward-compatible), tot .docx.
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    r = client.post(f"/api/dosar/{uid}/raport.docx", json=_payload())
    assert r.status_code == 200
    assert r.content[:2] == b"PK" and r.content[:4] != b"%PDF"   # .docx (zip OOXML), nu PDF
    r2 = client.post(f"/api/dosar/{uid}/raport.docx?fmt=pdf", json=_payload())
    assert r2.status_code == 200
    assert "application/pdf" not in r2.headers["content-type"]


def test_incarca_submis_marcheaza_asumat(client):
    # ADR-003 trigger #3 (decizia Adi #10): .docx submis → versiune „submis" + identitate asumată.
    import base64
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {"nume_client": "X"}}).json()["uuid"]
    payload = base64.b64encode(b"PK raport finalizat returnat de banca").decode()
    r = client.post(f"/api/dosar/{uid}/incarca-submis",
                    json={"nume_fisier": "raport-final.docx", "continut": payload})
    assert r.status_code == 200 and r.json()["asumat_la"]
    d = client.get(f"/api/dosar/{uid}").json()
    assert any(v["tip"] == "submis" for v in d.get("versiuni", []))


def test_deblocheaza_si_cloneaza_endpoints(client):
    # ADR-003: deblocheaza (motiv obligatoriu → 422 fără) + cloneaza (dosar nou); identitate read-only după asumare.
    import base64
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {"nume_client": "X", "au": "100"}}).json()["uuid"]
    client.post(f"/api/dosar/{uid}/incarca-submis",
                json={"nume_fisier": "r.docx", "continut": base64.b64encode(b"raport").decode()})
    # identitate înghețată după asumare (salveaza nu schimbă nume_client)
    client.post(f"/api/dosar/{uid}/salveaza", json={"nume_client": "ALTUL", "au": "999"})
    assert client.get(f"/api/dosar/{uid}").json()["wizard"]["nume_client"] == "X"
    # deblocare fără motiv → 422
    assert client.post(f"/api/dosar/{uid}/deblocheaza", json={"motiv": " "}).status_code == 422
    r = client.post(f"/api/dosar/{uid}/deblocheaza", json={"motiv": "typo nume"})
    assert r.status_code == 200 and r.json()["blocat"] is False
    # clonare → dosar nou diferit
    rc = client.post(f"/api/dosar/{uid}/cloneaza")
    assert rc.status_code == 200 and rc.json()["uuid"] != uid
    assert client.post("/api/dosar/zzz/cloneaza").status_code == 404


def test_dosar_html_arata_lock_dupa_asumare(client):
    # ADR-003 UI: pagina dosarului arată bannerul de lock + butoanele DOAR după asumare.
    import base64
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {"nume_client": "X"}}).json()["uuid"]
    # neasumat: fără banner (butonul de deblocare apare DOAR în blocul {% if blocat %})
    assert 'id="btn-delock"' not in client.get(f"/dosar/{uid}").text
    assert "var BLOCAT = false" in client.get(f"/dosar/{uid}").text
    client.post(f"/api/dosar/{uid}/incarca-submis",
                json={"nume_fisier": "r.docx", "continut": base64.b64encode(b"r").decode()})
    html = client.get(f"/dosar/{uid}").text
    assert 'id="btn-delock"' in html and 'id="btn-clone"' in html   # banner de lock prezent
    assert "var BLOCAT = true" in html


# ── Garda de VALOARE imposibila end-to-end (decizia owner 2026-06-10) ──────────────────────────────
# Audit advers: garda «pret corectat <=0 -> blocheaza» exista, dar valoarea negativa scapa la capete:
# /calcul o emitea (200 + valoare negativa), endpointul VECHI o persista + genera raport oficial pe ea,
# iar garda raportului nou cadea OPEN la exceptie. Aici fixam toate caile + protejam decizia I1.
def _payload_piata(adj_pct: str) -> dict:
    """Payload cu grila de piata (3 comparabile); primul are ajustarea de proprietate `adj_pct`."""
    p = _payload()
    p["metoda"] = "piata"
    p["comparables"] = [
        {"pret": "100000", "suprafata": "100",
         "adjustments": [{"element": "x", "tip": "procentuala", "valoare": adj_pct,
                          "etapa": "proprietate"}]},
        {"pret": "100000", "suprafata": "100"},
        {"pret": "100000", "suprafata": "100"},
    ]
    return p


def test_calcul_valoare_finala_negativa_blocata_422(client):
    # PROBLEMA 1: /calcul returna 200 + valoare_finala = −66666.67 (valoare imposibila, advisory).
    # Acum: 422 — un pret negativ nu e o estimare utila nici provizoriu.
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    r = client.post(f"/api/dosar/{uid}/calcul", json=_payload_piata("-5.0"))
    assert r.status_code == 422
    det = r.json()["detail"].lower()
    assert "imposibila" in det and ("valoarea finala" in det or "corectat" in det)


def test_calcul_pret_corectat_negativ_chiar_cu_finala_pozitiva_422(client):
    # Caz subtil: pret corectat <=0 dar media ramane POZITIVA (60000) — grila e poluata de un
    # comparabil nonsens. Carve-out-ul prinde si acest caz (decizia owner: «pret corectat <=0»).
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    r = client.post(f"/api/dosar/{uid}/calcul", json=_payload_piata("-1.20"))
    assert r.status_code == 422 and "corectat" in r.json()["detail"].lower()


def test_calcul_blocaj_de_date_ramane_advisory_200(client):
    # REGRESIE decizia re-audit I1: un blocaj de DATE (Au>Acd) cu valoare POZITIVA ramane advisory
    # in /calcul (200 + alerta) — carve-out-ul NU trebuie sa-l promoveze la 422. Protejeaza granita.
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    p = _payload()
    p["building"]["au"] = "200"                        # Au(200) > Acd(120) -> blocheaza de DATE
    r = client.post(f"/api/dosar/{uid}/calcul", json=p)
    assert r.status_code == 200
    assert any(a["nivel"] == "blocheaza" for a in r.json()["alerte"])   # vizibila ca alerta, nu 422


def test_raport_nou_blocat_pe_valoare_imposibila_422(client):
    # Documentul oficial nou e blocat pe valoare imposibila (deja prin pret corectat <=0, acum si
    # prin validarea incrucisata «valoare finala <=0» — vezi fix-ul de completitudine).
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    r = client.post(f"/api/dosar/{uid}/raport.docx", json=_payload_piata("-5.0"))
    assert r.status_code == 422 and "blocat" in r.json()["detail"].lower()


def test_raport_nou_fail_closed_daca_validarea_arunca(client, monkeypatch):
    # PROBLEMA 3 (fail-open): codul vechi prindea exceptia din valideaza() si seta blocante=[] -> raport
    # GENERAT pe validare esuata. Acum garda e FAIL-CLOSED: o exceptie -> 422, nu document mut.
    def _arunca(*a, **k):
        raise ArithmeticError("validare degenerata simulata")
    monkeypatch.setattr("evaluare.web.routers.curent.valideaza", _arunca)
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    r = client.post(f"/api/dosar/{uid}/raport.docx", json=_payload())   # input altfel valid
    assert r.status_code == 422 and "fail-closed" in r.json()["detail"].lower()


def test_evaluare_veche_nu_persista_valoare_imposibila_422(client):
    # PROBLEMA 2: endpointul VECHI /api/evaluare persista dosarul cu valoare_finala = −66666.67.
    # Acum: 422 inainte de save -> dosarul NU se creeaza (nicio scriere orfana cu valoare negativa).
    r = client.post("/api/evaluare", json=_payload_piata("-5.0"))
    assert r.status_code == 422 and "imposibila" in r.json()["detail"].lower()
    assert client.get("/api/evaluare/1").status_code == 404   # nimic persistat


def test_raport_vechi_blocat_pe_dosar_legacy_422(client, monkeypatch):
    # PROBLEMA 2 (aparare legacy): un dosar deja persistat cu valoare imposibila (inainte de fix) NU
    # mai produce raport oficial. Simulam un dosar legacy dezactivand garda DOAR la create (ca vechiul
    # cod), apoi cerem raportul -> 422 (garda raportului vechi, pe context, blocheaza).
    monkeypatch.setattr("evaluare.web.routers.evaluare.valoare_imposibila", lambda ctx: [])
    eid = client.post("/api/evaluare", json=_payload_piata("-5.0")).json()["id"]   # persistat (legacy)
    monkeypatch.undo()
    r = client.get(f"/api/evaluare/{eid}/raport.docx")
    assert r.status_code == 422 and "blocat" in r.json()["detail"].lower()


def test_lock_unlock_concurenta_endpoints(client):
    # ADR-003 item 7: lock detectează altă fereastră (alt token); unlock eliberează; 404 pe inexistent.
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {"nume_client": "X"}}).json()["uuid"]
    assert client.post(f"/api/dosar/{uid}/lock", json={"token": "A"}).json()["concurent"] is False
    assert client.post(f"/api/dosar/{uid}/lock", json={"token": "B"}).json()["concurent"] is True
    assert client.post(f"/api/dosar/{uid}/unlock", json={"token": "B"}).json()["ok"] is True
    assert client.post(f"/api/dosar/{uid}/lock", json={"token": "C"}).json()["concurent"] is False
    assert client.post("/api/dosar/zzz/lock", json={"token": "A"}).status_code == 404


def test_dosar_json_corupt_nu_crapa(client):
    # robustețe: dosar.json corupt -> 404 clar (nu 500); workspace-ul nu trebuie să crape la deschidere.
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {"nume_client": "X"}}).json()["uuid"]
    (client._baza / "date" / "dosare" / uid / "dosar.json").write_text("{ corupt", encoding="utf-8")
    assert client.get(f"/dosar/{uid}").status_code == 404
    assert client.get(f"/api/dosar/{uid}").status_code == 404


def test_calcul_alerte_sunt_expuse(client):
    # contractul „aplicația avertizează, nu decide": Au>Acd trebuie să apară în `alerte` (nu 200 mut).
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    p = _payload()
    p["building"]["au"] = "200"          # Au(200) > Acd(120) -> alertă de proprietate
    r = client.post(f"/api/dosar/{uid}/calcul", json=p)
    assert r.status_code == 200
    alerte = r.json()["alerte"]
    assert isinstance(alerte, list) and len(alerte) >= 1


def test_import_prea_mare_413(client):
    _cont(client)
    big = "A" * 35_000_001            # peste limita anti-DoS
    assert client.post("/api/dosar/import-docx",
                       json={"nume_fisier": "x.docx", "continut": big}).status_code == 413
    assert client.post("/api/ingestie", json={"tip": "cf", "continut": big}).status_code == 413


def test_host_nelocal_respins_403(client):
    # gardă anti DNS-rebinding: Host ne-local -> 403; default (testserver) -> OK.
    assert client.get("/incepe", headers={"host": "evil.com"}).status_code == 403
    assert client.get("/incepe").status_code == 200


def test_ssrf_url_guard():
    from evaluare.importers.url_parser import _url_public_sigur
    assert _url_public_sigur("http://127.0.0.1:5432") is False      # loopback
    assert _url_public_sigur("http://169.254.169.254/") is False    # link-local (metadata)
    assert _url_public_sigur("http://10.0.0.5/x") is False          # privat
    assert _url_public_sigur("file:///etc/passwd") is False         # schemă nepermisă


def test_import_url_gol_422_nu_500(client):
    # robustețe (audit D schemathesis pe live): url="" ajungea la parser -> 500 nehandelat.
    # Garda SSRF respinge URL-ul -> ValueError prins în router -> 422 clar (nu 500).
    r = client.post("/api/import-url", json={"url": ""})
    assert r.status_code == 422 and "url" in r.json()["detail"].lower()


def test_fetch_html_blocheaza_redirect_spre_intern(monkeypatch):
    # SSRF prin redirect: un host PUBLIC raspunde 302 -> adresa INTERNA. Garda re-valideaza FIECARE
    # Location si respinge (nu urmeaza orbeste). Vezi fix-ul allow_redirects=False + bucla manuala.
    import evaluare.importers.url_parser as up

    class FakeResp:
        is_redirect = True
        headers = {"Location": "http://169.254.169.254/latest/meta-data/"}

        def raise_for_status(self):
            pass

    monkeypatch.setattr(up.requests, "get", lambda *a, **k: FakeResp())
    try:
        up.fetch_html("http://93.184.216.34/")              # IP public valid -> 302 -> intern
        raise AssertionError("ar fi trebuit sa respinga redirect-ul spre adresa interna")
    except ValueError as e:
        assert "SSRF" in str(e)


def test_comparabile_resping_suprafata_si_pret_zero():
    # audit B: comparabil cu suprafata/pret 0 (sau negativ) -> respins la validare (Field gt=0), nu 500
    # din DivisionByZero in market.py / pret_unitar_brut.
    from pydantic import ValidationError

    from evaluare.models.comparable import Comparable, LandComparable, RentComparable
    for kw in ({"pret": 100, "suprafata": 0}, {"pret": 0, "suprafata": 100}):
        with pytest.raises(ValidationError):
            Comparable(**kw)
    with pytest.raises(ValidationError):
        LandComparable(pret_mp=0, suprafata=100)
    with pytest.raises(ValidationError):
        RentComparable(chirie_mp=10, suprafata=0)


def test_security_headers_prezente(client):
    # audit C/D (defense-in-depth): nosniff + anti-clickjacking + CSP + referrer + permissions.
    r = client.get("/incepe")
    assert r.headers.get("X-Content-Type-Options") == "nosniff"
    assert r.headers.get("X-Frame-Options") == "DENY"
    assert "Content-Security-Policy" in r.headers
    assert "Referrer-Policy" in r.headers
    assert "Permissions-Policy" in r.headers


def test_cnp_redaction_prefix_9():
    from evaluare.ai import narrative
    rx = narrative._PII_REZIDUAL[0][0]
    assert rx.search("9800101123456")    # rezident străin (prefix 9) — acum prins
    assert rx.search("1800101123456")    # cetățean (prefix 1) — în continuare prins


def test_dosar_gating_descopera_markup_prezent(client):
    # Gating «Descoperă» (sub-tabul Comparabile): pagina dosarului include atenționarea de gating
    # + afișajul fix al zonei + paznicul JS care propagă județ+localitate și le îngheață read-only.
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {"nume_client": "X"}}).json()["uuid"]
    html = client.get(f"/dosar/{uid}").text
    # atenționarea cu mesajul exact cerut de Adi
    assert 'id="gate-descopera"' in html
    assert "Completează județul și localitatea pe tabul Proprietate mai întâi." in html
    # afișajul fix (read-only) al zonei propagate pe Descoperă
    assert 'id="d-zona-fix"' in html and 'id="d-zona-judet"' in html and 'id="d-zona-localitate"' in html
    # logica de gating: paznicul + funcția care verifică dacă ambele câmpuri sunt completate
    assert "propGata" in html and "inghetZona" in html


# ── Candidați salvați din «Descoperă» (în așteptare per dosar) ───────────────────
def _candidat_api(url="https://imobiliare.ro/anunt/1", **extra):
    c = {"url": url, "titlu": "Casa P+1", "pret": "120000", "suprafata": "100",
         "teren": "500", "pret_mp": "1200", "poza": "https://x.ro/p.jpg",
         "relevanta": 0.87, "localitate": "Sinaia", "distanta_km": 3.2}
    c.update(extra)
    return c


def test_candidat_salvat_salveaza_listeaza_sterge(client):
    # Contract complet: POST salvează (dict întreg), GET listează, POST .../sterge scoate pe url.
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {"nume_client": "X"}}).json()["uuid"]
    assert client.get(f"/api/dosar/{uid}/candidati-salvati").json()["candidati"] == []
    r = client.post(f"/api/dosar/{uid}/candidat-salvat", json=_candidat_api())
    assert r.status_code == 200
    body = r.json()
    assert body["count"] == 1 and body["candidati"][0]["pret_mp"] == "1200"
    # GET reflectă starea
    assert len(client.get(f"/api/dosar/{uid}/candidati-salvati").json()["candidati"]) == 1
    # al doilea url
    client.post(f"/api/dosar/{uid}/candidat-salvat", json=_candidat_api(url="https://imobiliare.ro/anunt/2"))
    assert client.post(f"/api/dosar/{uid}/candidat-salvat", json=_candidat_api()).json()["count"] == 2  # dedup
    # ștergere pe url
    r2 = client.post(f"/api/dosar/{uid}/candidat-salvat/sterge",
                     json={"url": "https://imobiliare.ro/anunt/1"})
    assert r2.status_code == 200
    assert [c["url"] for c in r2.json()["candidati"]] == ["https://imobiliare.ro/anunt/2"]


def test_candidat_salvat_fara_url_422(client):
    # input gardat: candidat fără url -> 422 (nu 500).
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    assert client.post(f"/api/dosar/{uid}/candidat-salvat",
                       json={"titlu": "fără url"}).status_code == 422
    assert client.post(f"/api/dosar/{uid}/candidat-salvat",
                       json={"url": "  "}).status_code == 422


def test_candidat_salvat_uid_invalid_404(client):
    # gard UUID (anti path-traversal): uid non-UUID -> 404 pe toate cele trei rute.
    assert client.get("/api/dosar/not-a-uuid/candidati-salvati").status_code == 404
    assert client.post("/api/dosar/not-a-uuid/candidat-salvat",
                       json=_candidat_api()).status_code == 404
    assert client.post("/api/dosar/not-a-uuid/candidat-salvat/sterge",
                       json={"url": "https://x.ro/a"}).status_code == 404


def test_candidat_salvat_dosar_inexistent_404(client):
    # uid UUID-valid dar dosar inexistent -> 404 (nu folder fantomă).
    import uuid as _uuid
    _cont(client)
    fake = str(_uuid.uuid4())
    assert client.post(f"/api/dosar/{fake}/candidat-salvat",
                       json=_candidat_api()).status_code == 404


def test_dosar_inexistent_404(client):
    assert client.get("/dosar/nope").status_code == 404
    assert client.get("/api/dosar/nope").status_code == 404


def test_sterge_dosar(client):
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    client.post(f"/api/dosar/{uid}/sterge")
    assert client.get(f"/api/dosar/{uid}").status_code == 404


def test_scoate_dosar_disparut_din_index(client):
    import shutil

    from evaluare import dosare_fs as fs
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {"nume_client": "X"}}).json()["uuid"]
    client.get("/incepe")                                  # intră în index „ultima vedere"
    shutil.rmtree(client._baza / "date" / "dosare" / uid)  # folderul dispare de pe disc
    assert uid in fs._citeste_index()                      # încă în index (dispărut)
    assert client.post(f"/api/dosar/{uid}/scoate-din-index").json()["ok"] is True
    assert uid not in fs._citeste_index()                  # scos din index


# ── Import .docx ─────────────────────────────────────────────────────────────
def _docx_b64(tmp_path, nume):
    import base64

    import docx
    doc = docx.Document()
    doc.add_paragraph("Utilizarea desemnata a evaluării este garantarea împrumutului.")
    f = tmp_path / nume
    doc.save(str(f))
    return "data:application/octet-stream;base64," + base64.b64encode(f.read_bytes()).decode()


def test_import_docx_creeaza_dosar_precompletat(client, tmp_path):
    _cont(client)
    b64 = _docx_b64(tmp_path, "555 Maria Ionescu casa Sinaia.docx")
    r = client.post("/api/dosar/import-docx",
                    json={"nume_fisier": "555 Maria Ionescu casa Sinaia.docx", "continut": b64})
    assert r.status_code == 200
    d = r.json()
    assert d["wizard"]["id_client"] == "555"
    assert d["wizard"]["nume_client"] == "Maria Ionescu"
    assert d["wizard"]["judet"] == "Prahova"
    # raportul sursă e atașat ca versiune
    folder = client._baza / "date" / "dosare" / d["uuid"]
    assert len(list(folder.glob("raport-*.docx"))) == 1


def test_import_docx_necesita_cont(client, tmp_path):
    b64 = _docx_b64(tmp_path, "1 X casa Cluj.docx")
    r = client.post("/api/dosar/import-docx", json={"nume_fisier": "1 X casa Cluj.docx", "continut": b64})
    assert r.status_code == 403


def test_import_docx_continut_invalid_400(client):
    _cont(client)
    r = client.post("/api/dosar/import-docx",
                    json={"nume_fisier": "x.docx", "continut": "@@@nu e base64@@@"})
    assert r.status_code == 400


# ── #7 coverage cod nou: /setari, /api/cont edit, recalc nume API, identitate raport ──────────
def test_setari_get_cu_si_fara_cont(client):
    # /setari (pagina noua de vizualizare/editare cont) — accesibila si fara cont
    assert client.get("/setari").status_code == 200
    _cont(client)
    r = client.get("/setari")
    assert r.status_code == 200 and "Adi S" in r.text and "8717" in r.text


def test_api_cont_edit_actualizeaza(client):
    # /api/cont re-postat = EDIT: actualizeaza identitatea evaluatorului persistata
    _cont(client)
    assert client.post("/api/cont", json={"nume": "Adi Schimbat",
                                          "legitimatie": "9999"}).status_code == 200
    pagina = client.get("/cont").text
    assert "Adi Schimbat" in pagina and "9999" in pagina and "8717" not in pagina


def test_api_dosar_salveaza_recalc_nume_end_to_end(client):
    # end-to-end: dosar creat GOL -> nume cu „?"; dupa /salveaza cu identitate -> fara „?" (bug Adi)
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    assert "?" in client.get(f"/api/dosar/{uid}").json()["nume"]
    client.post(f"/api/dosar/{uid}/salveaza",
                json={"id_client": "D9", "nume_client": "Pop", "tip_proprietate": "casa"})
    assert "?" not in client.get(f"/api/dosar/{uid}").json()["nume"]


def test_raport_contine_identitatea_lucrarii(client):
    # #7: identitatea lucrarii (scopul evaluarii) apare in raportul .docx generat
    import io

    import docx
    _cont(client)
    uid = client.post("/api/dosar", json={"wizard": {}}).json()["uuid"]
    p = _payload()
    p["meta"]["scop"] = "garantarea imprumutului bancar XYZ"
    r = client.post(f"/api/dosar/{uid}/raport.docx", json=p)
    assert r.status_code == 200
    text = "\n".join(par.text for par in docx.Document(io.BytesIO(r.content)).paragraphs)
    assert "garantarea imprumutului bancar XYZ" in text
