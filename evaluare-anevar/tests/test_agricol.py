from decimal import Decimal

from docx import Document

from evaluare.models.meta import EvaluationMeta
from evaluare.models.narrative import NarrativeSection
from evaluare.models.property import BuildingData, CostElement, DepreciationPoint, LandData
from evaluare.models.report_context import ReportContext
from evaluare.models.results import CostResult, ReconciledResult
from evaluare.profil import AGRICOL
from evaluare.report.generator import genereaza_raport


def test_profil_agricol():
    assert AGRICOL.tip_activ == "agricol"
    assert AGRICOL.abordari_aplicabile == ["comparatie"]
    assert AGRICOL.ghid == "GEV_630"


def test_land_campuri_agricole_optionale():
    land = LandData(suprafata=Decimal("5000"), categorie_folosinta="arabil", clasa_calitate=2)
    assert land.categorie_folosinta == "arabil" and land.clasa_calitate == 2
    l2 = LandData(suprafata=Decimal("500"))
    assert l2.categorie_folosinta is None


def _ctx_agricol() -> ReportContext:
    meta = EvaluationMeta(
        client_nume="Ion Popescu", adresa="Str. Exemplu 1, Bucuresti",
        numar_cadastral="123456", carte_funciara="CF123456",
        evaluator_nume="Maria Ionescu", evaluator_legitimatie="19567",
        data_evaluarii="2026-01-16", data_raportului="2026-01-16",
    )
    return ReportContext(
        meta=meta,
        land=LandData(suprafata=Decimal("5000"), categorie="extravilan"),
        building=BuildingData(
            au=Decimal("0"), acd=Decimal("0"), an_referinta=2025,
            elements=[
                CostElement(element="Structura", cod="X", um="mp",
                            cantitate=Decimal("1"), cost_unitar=Decimal("1"),
                            an_pif=2015),
            ],
            depreciation_points=[DepreciationPoint(varsta=10, depreciere=Decimal("0.1"))],
        ),
        cost_result=CostResult(cib=Decimal("1"), vcp=Decimal("1"),
                               depreciere_fizica=Decimal("0.10"), cin=Decimal("1"),
                               valoare_cost=Decimal("50000")),
        reconciled=ReconciledResult(valoare_finala=Decimal("50000"),
                                    metoda_selectata="piata"),
        narrative=[NarrativeSection(capitol="Analiza celei mai bune utilizari (CMBU)",
                                    text="Cea mai buna utilizare este agricola.")],
    )


def _all_text(path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_raport_cu_teren_agricol(tmp_path):
    ctx = _ctx_agricol()
    ctx.land.categorie_folosinta = "arabil"
    ctx.land.clasa_calitate = 2
    out = tmp_path / "raport_agricol.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Teren agricol" in text
    assert "arabil" in text
    assert "clasa de calitate 2" in text


def test_raport_fara_teren_agricol(tmp_path):
    ctx = _ctx_agricol()
    # categorie_folosinta is None by default — no agricultural line expected
    out = tmp_path / "raport_neagricol.docx"
    genereaza_raport(ctx, out)
    text = _all_text(out)
    assert "Teren agricol" not in text
